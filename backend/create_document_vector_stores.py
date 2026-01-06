"""
Create FAISS vector stores for document agent test documents.

This script builds embeddings for each test document and saves FAISS indices.
Required for RAG-based document queries to work.
"""

import sys
from pathlib import Path
import logging

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Get workspace root
WORKSPACE_ROOT = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(WORKSPACE_ROOT / "backend"))

def create_vector_stores():
    """Create FAISS vector stores for all test documents."""
    
    try:
        from langchain_community.vectorstores import FAISS
        from langchain_huggingface import HuggingFaceEmbeddings
        from langchain_community.document_loaders import PyPDFLoader
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from docx import Document as DocxDocument
        from langchain.schema import Document
    except ImportError as e:
        logger.error(f"Missing required packages: {e}")
        logger.info("Install with: pip install langchain langchain-community langchain-huggingface python-docx pypdf")
        return False
    
    # Create storage directory
    vector_store_dir = WORKSPACE_ROOT / "storage" / "vector_store"
    vector_store_dir.mkdir(parents=True, exist_ok=True)
    
    # Initialize embeddings (same as document agent)
    logger.info("Loading embeddings model: all-mpnet-base-v2")
    embeddings = HuggingFaceEmbeddings(model_name='all-mpnet-base-v2')
    
    # Text splitter for chunking documents
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", " ", ""]
    )
    
    # Document configuration
    documents_config = {
        "phd_thesis": {
            "path": WORKSPACE_ROOT / "backend" / "tests" / "document_agent" / "test_data" / "2004_phdthesis_lip6.pdf",
            "type": "pdf",
            "store_name": "phd_thesis_store"
        },
        "research_paper": {
            "path": WORKSPACE_ROOT / "backend" / "tests" / "document_agent" / "test_data" / "2212.07286v2.pdf",
            "type": "pdf",
            "store_name": "research_paper_store"
        },
        "sales_document": {
            "path": WORKSPACE_ROOT / "backend" / "tests" / "document_agent" / "test_data" / "SampleDocs-sales-sample-data.docx",
            "type": "docx",
            "store_name": "sales_document_store"
        },
        "generic_document": {
            "path": WORKSPACE_ROOT / "backend" / "tests" / "document_agent" / "test_data" / "SampleDocs-Test Word File With Dummy Data.docx",
            "type": "docx",
            "store_name": "generic_document_store"
        }
    }
    
    # Create vector stores for each document
    for doc_key, config in documents_config.items():
        doc_path = config["path"]
        store_name = config["store_name"]
        store_path = vector_store_dir / store_name
        
        # Check if file exists
        if not doc_path.exists():
            logger.error(f"Document not found: {doc_path}")
            continue
        
        # Skip if vector store already exists
        if store_path.exists():
            logger.info(f"[SKIP] Vector store already exists: {store_name}")
            continue
        
        logger.info(f"[START] Creating vector store: {store_name}")
        
        try:
            # Load documents based on type
            if config["type"] == "pdf":
                logger.info(f"  - Loading PDF: {doc_path.name}")
                loader = PyPDFLoader(str(doc_path))
                docs = loader.load()
                logger.info(f"  - Loaded {len(docs)} pages")
            
            elif config["type"] == "docx":
                logger.info(f"  - Loading DOCX: {doc_path.name}")
                docx_doc = DocxDocument(doc_path)
                docs = []
                for idx, para in enumerate(docx_doc.paragraphs):
                    if para.text.strip():
                        docs.append(Document(
                            page_content=para.text,
                            metadata={"source": str(doc_path), "paragraph": idx}
                        ))
                logger.info(f"  - Loaded {len(docs)} paragraphs")
            
            # Split documents into chunks
            logger.info(f"  - Splitting into chunks...")
            chunks = text_splitter.split_documents(docs)
            logger.info(f"  - Created {len(chunks)} chunks")
            
            if not chunks:
                logger.warning(f"  - No chunks created for {doc_key}")
                continue
            
            # Create FAISS vector store
            logger.info(f"  - Building embeddings...")
            vector_store = FAISS.from_documents(chunks, embeddings)
            logger.info(f"  - Creating FAISS index...")
            
            # Save vector store
            vector_store.save_local(str(store_path))
            logger.info(f"[OK] Vector store created: {store_name}")
            logger.info(f"     Location: {store_path}")
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to create vector store for {doc_key}: {e}")
            import traceback
            traceback.print_exc()
            continue
    
    logger.info("\n" + "="*70)
    logger.info("Vector store creation complete!")
    logger.info("="*70)
    return True


if __name__ == "__main__":
    success = create_vector_stores()
    sys.exit(0 if success else 1)
