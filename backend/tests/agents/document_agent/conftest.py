"""
Pytest configuration and shared fixtures for document agent tests.
"""

import pytest
import os
import tempfile
import shutil
from pathlib import Path
from typing import Generator
from io import BytesIO

# Test data imports
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib.pagesizes import letter

# Agent imports (will be mocked where needed)
import sys
backend_path = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(backend_path))


# ============== DIRECTORY FIXTURES ==============

@pytest.fixture(scope="session")
def test_data_dir() -> Path:
    """Base directory for test data files."""
    base_dir = Path(__file__).parent / "test_data"
    base_dir.mkdir(exist_ok=True)
    return base_dir


@pytest.fixture
def temp_storage(tmp_path) -> Generator[Path, None, None]:
    """Temporary storage directory for test files."""
    storage_dir = tmp_path / "document_storage"
    storage_dir.mkdir()
    yield storage_dir
    # Cleanup happens automatically with tmp_path


# ============== PDF FIXTURES ==============

@pytest.fixture
def simple_pdf(test_data_dir: Path) -> Path:
    """Generate a simple single-page PDF with text."""
    pdf_path = test_data_dir / "simple_test.pdf"
    
    if not pdf_path.exists():
        c = pdf_canvas.Canvas(str(pdf_path), pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(100, 750, "Test Document")
        c.drawString(100, 730, "This is a simple test PDF with text content.")
        c.drawString(100, 710, "Line 3 of text.")
        c.save()
    
    return pdf_path


@pytest.fixture
def complex_pdf(test_data_dir: Path) -> Path:
    """Generate a multi-page PDF with varied content."""
    pdf_path = test_data_dir / "complex_test.pdf"
    
    if not pdf_path.exists():
        c = pdf_canvas.Canvas(str(pdf_path), pagesize=letter)
        
        # Page 1 - Text
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 750, "Complex Document - Page 1")
        c.setFont("Helvetica", 12)
        for i in range(10):
            c.drawString(100, 700 - i * 20, f"Line {i+1} with various content types.")
        
        c.showPage()
        
        # Page 2 - More text
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 750, "Page 2 - Continued")
        c.setFont("Helvetica", 12)
        c.drawString(100, 720, "Additional content on second page.")
        
        c.save()
    
    return pdf_path


@pytest.fixture
def empty_pdf(test_data_dir: Path) -> Path:
    """Generate an empty PDF."""
    pdf_path = test_data_dir / "empty_test.pdf"
    
    if not pdf_path.exists():
        c = pdf_canvas.Canvas(str(pdf_path), pagesize=letter)
        c.showPage()
        c.save()
    
    return pdf_path


# ============== DOCX FIXTURES ==============

@pytest.fixture
def simple_docx(test_data_dir: Path) -> Path:
    """Generate a simple DOCX with text."""
    docx_path = test_data_dir / "simple_test.docx"
    
    if not docx_path.exists():
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a simple test DOCX file.')
        doc.add_paragraph('It contains multiple paragraphs.')
        doc.add_paragraph('And basic formatting.')
        doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def complex_docx(test_data_dir: Path) -> Path:
    """Generate a complex DOCX with tables, lists, and formatting."""
    docx_path = test_data_dir / "complex_test.docx"
    
    if not docx_path.exists():
        doc = Document()
        
        # Title
        doc.add_heading('Complex Test Document', 0)
        
        # Text with formatting
        p = doc.add_paragraph('This document contains ')
        p.add_run('bold text').bold = True
        p.add_run(' and ')
        p.add_run('italic text').italic = True
        p.add_run('.')
        
        # Bulleted list
        doc.add_paragraph('Item 1', style='List Bullet')
        doc.add_paragraph('Item 2', style='List Bullet')
        doc.add_paragraph('Item 3', style='List Bullet')
        
        # Numbered list
        doc.add_paragraph('First', style='List Number')
        doc.add_paragraph('Second', style='List Number')
        doc.add_paragraph('Third', style='List Number')
        
        # Table
        table = doc.add_table(rows=3, cols=3)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f'Cell {i},{j}'
        
        # Section heading
        doc.add_heading('Section 2', 2)
        doc.add_paragraph('Additional content in second section.')
        
        doc.save(str(docx_path))
    
    return docx_path


@pytest.fixture
def empty_docx(test_data_dir: Path) -> Path:
    """Generate an empty DOCX."""
    docx_path = test_data_dir / "empty_test.docx"
    
    if not docx_path.exists():
        doc = Document()
        doc.save(str(docx_path))
    
    return docx_path


# ============== IMAGE FIXTURES ==============

@pytest.fixture
def simple_image(test_data_dir: Path) -> Path:
    """Generate a simple test image."""
    img_path = test_data_dir / "simple_test.png"
    
    if not img_path.exists():
        # Create a 800x600 white image with some colored rectangles
        img = Image.new('RGB', (800, 600), color='white')
        pixels = img.load()
        
        # Add some colored shapes
        for i in range(200, 400):
            for j in range(200, 400):
                pixels[i, j] = (255, 0, 0)  # Red square
        
        for i in range(450, 650):
            for j in range(200, 400):
                pixels[i, j] = (0, 0, 255)  # Blue square
        
        img.save(str(img_path))
    
    return img_path


@pytest.fixture
def high_res_image(test_data_dir: Path) -> Path:
    """Generate a high-resolution test image."""
    img_path = test_data_dir / "highres_test.png"
    
    if not img_path.exists():
        # Create a 3000x2000 image
        img = Image.new('RGB', (3000, 2000), color='lightblue')
        img.save(str(img_path))
    
    return img_path


@pytest.fixture
def jpg_image(test_data_dir: Path) -> Path:
    """Generate a JPG image."""
    img_path = test_data_dir / "test_image.jpg"
    
    if not img_path.exists():
        img = Image.new('RGB', (640, 480), color='green')
        img.save(str(img_path), 'JPEG')
    
    return img_path


# ============== TEXT FIXTURES ==============

@pytest.fixture
def simple_txt(test_data_dir: Path) -> Path:
    """Generate a simple text file."""
    txt_path = test_data_dir / "simple_test.txt"
    
    if not txt_path.exists():
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("This is a simple text file.\n")
            f.write("It has multiple lines.\n")
            f.write("Line 3 of content.\n")
    
    return txt_path


# ============== INVALID FILE FIXTURES ==============

@pytest.fixture
def corrupted_pdf(test_data_dir: Path) -> Path:
    """Generate a corrupted PDF file."""
    pdf_path = test_data_dir / "corrupted.pdf"
    
    if not pdf_path.exists():
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4\n%corrupted data\nthis is not valid PDF')
    
    return pdf_path


@pytest.fixture
def unsupported_file(test_data_dir: Path) -> Path:
    """Generate an unsupported file type."""
    file_path = test_data_dir / "unsupported.xyz"
    
    if not file_path.exists():
        with open(file_path, 'w') as f:
            f.write("This is an unsupported file format")
    
    return file_path


@pytest.fixture
def large_file(test_data_dir: Path) -> Path:
    """Generate a file exceeding size limit (100MB)."""
    file_path = test_data_dir / "too_large.txt"
    
    # Only create if doesn't exist (it's large)
    if not file_path.exists():
        with open(file_path, 'wb') as f:
            f.write(b'x' * (101 * 1024 * 1024))  # 101 MB
    
    return file_path


# ============== AGENT FIXTURES ==============

@pytest.fixture
def document_agent():
    """Initialize a document agent instance."""
    from backend.agents.document_agent import DocumentAgent
    
    agent = DocumentAgent()
    yield agent
    
    # Cleanup
    if hasattr(agent, 'cleanup'):
        agent.cleanup()


@pytest.fixture
def mock_llm():
    """Mock LLM for testing without API calls."""
    from unittest.mock import Mock
    
    llm = Mock()
    llm.predict.return_value = "This is a mocked LLM response."
    
    return llm


# ============== SESSION FIXTURES ==============

@pytest.fixture
def test_session(temp_storage):
    """Create a test session."""
    session_data = {
        'session_id': 'test-session-123',
        'storage_dir': str(temp_storage),
        'files': {},
        'history': []
    }
    
    return session_data


@pytest.fixture
def populated_session(test_session, simple_docx):
    """Session with a document already uploaded."""
    test_session['files']['test-file-1'] = {
        'file_id': 'test-file-1',
        'filename': 'simple_test.docx',
        'filepath': str(simple_docx),
        'file_type': 'docx',
        'size': simple_docx.stat().st_size
    }
    
    return test_session


# ============== UTILITY FIXTURES ==============

@pytest.fixture
def file_upload_data(simple_pdf):
    """Simulated file upload data."""
    with open(simple_pdf, 'rb') as f:
        content = f.read()
    
    return {
        'filename': 'test_upload.pdf',
        'content': content,
        'content_type': 'application/pdf'
    }


@pytest.fixture
def cleanup_test_files():
    """Cleanup fixture that runs after tests."""
    yield
    # Add any additional cleanup logic here


# ============== PYTEST CONFIGURATION ==============

def pytest_configure(config):
    """Configure pytest with custom markers."""
    config.addinivalue_line(
        "markers", "slow: marks tests as slow (deselect with '-m \"not slow\"')"
    )
    config.addinivalue_line(
        "markers", "integration: marks tests as integration tests"
    )
    config.addinivalue_line(
        "markers", "unit: marks tests as unit tests"
    )
    config.addinivalue_line(
        "markers", "performance: marks tests as performance tests"
    )


# ============== PARAMETRIZE HELPERS ==============

SUPPORTED_FORMATS = ['pdf', 'docx', 'txt', 'png', 'jpg']
UNSUPPORTED_FORMATS = ['xlsx', 'zip', 'exe', 'mp4']

@pytest.fixture(params=SUPPORTED_FORMATS)
def supported_format(request):
    """Parametrize over supported file formats."""
    return request.param


@pytest.fixture(params=UNSUPPORTED_FORMATS)
def unsupported_format(request):
    """Parametrize over unsupported file formats."""
    return request.param
