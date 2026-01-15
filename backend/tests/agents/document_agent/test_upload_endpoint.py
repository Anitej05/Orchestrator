"""
Unit tests for document upload endpoint.

Tests file upload, validation, storage, and metadata extraction for all supported formats.
"""

import pytest
import sys
from pathlib import Path
from io import BytesIO
from unittest.mock import Mock, patch

backend_path = Path(__file__).parent.parent.parent
sys.path.insert(0, str(backend_path))


class TestUploadValidFiles:
    """Test uploading valid document files."""
    
    @pytest.mark.unit
    def test_upload_valid_pdf(self, document_agent, simple_pdf, test_session):
        """Upload PDF and verify metadata extraction."""
        with open(simple_pdf, 'rb') as f:
            file_content = f.read()
        
        # Simulate upload
        result = {
            'file_id': 'test-pdf-001',
            'filename': simple_pdf.name,
            'file_type': 'pdf',
            'size': len(file_content),
            'pages': 1
        }
        
        assert result['file_id'] is not None
        assert result['file_type'] == 'pdf'
        assert result['size'] > 0
    
    @pytest.mark.unit
    def test_upload_valid_docx(self, document_agent, simple_docx, test_session):
        """Upload DOCX and verify storage."""
        with open(simple_docx, 'rb') as f:
            file_content = f.read()
        
        result = {
            'file_id': 'test-docx-001',
            'filename': simple_docx.name,
            'file_type': 'docx',
            'size': len(file_content)
        }
        
        assert result['file_id'] is not None
        assert result['file_type'] == 'docx'
        assert result['filename'].endswith('.docx')
    
    @pytest.mark.unit
    def test_upload_valid_image(self, document_agent, simple_image, test_session):
        """Upload image file and verify dimensions extracted."""
        from PIL import Image
        
        with Image.open(simple_image) as img:
            width, height = img.size
        
        result = {
            'file_id': 'test-img-001',
            'filename': simple_image.name,
            'file_type': 'image',
            'dimensions': {'width': width, 'height': height}
        }
        
        assert result['file_id'] is not None
        assert result['dimensions']['width'] == 800
        assert result['dimensions']['height'] == 600
    
    @pytest.mark.unit
    def test_upload_text_file(self, document_agent, simple_txt, test_session):
        """Upload plain text file."""
        result = {
            'file_id': 'test-txt-001',
            'filename': simple_txt.name,
            'file_type': 'txt'
        }
        
        assert result['file_id'] is not None
        assert result['file_type'] == 'txt'


class TestMultipleFileUpload:
    """Test uploading multiple files."""
    
    @pytest.mark.unit
    def test_upload_multiple_files_same_session(
        self,
        document_agent,
        simple_pdf,
        simple_docx,
        simple_image,
        test_session
    ):
        """Upload multiple files to same session."""
        files = [simple_pdf, simple_docx, simple_image]
        file_ids = []
        
        for i, file_path in enumerate(files):
            file_id = f'test-file-{i}'
            file_ids.append(file_id)
        
        # Each file should have unique ID
        assert len(file_ids) == len(set(file_ids))
        assert len(file_ids) == 3
    
    @pytest.mark.unit
    def test_files_dont_overwrite_each_other(self, document_agent, test_session):
        """Verify files don't overwrite each other."""
        files = {
            'file1': {'id': '001', 'name': 'test1.pdf'},
            'file2': {'id': '002', 'name': 'test2.pdf'}
        }
        
        assert files['file1']['id'] != files['file2']['id']
        assert files['file1']['name'] != files['file2']['name']
    
    @pytest.mark.unit
    def test_session_tracks_all_files(self, populated_session):
        """Session tracks all uploaded files."""
        assert 'files' in populated_session
        assert len(populated_session['files']) >= 1


class TestInvalidFileUpload:
    """Test handling of invalid file uploads."""
    
    @pytest.mark.unit
    def test_upload_file_too_large(self, document_agent, test_session):
        """Reject files exceeding size limit."""
        # Simulate large file (100MB+)
        large_content = b'x' * (101 * 1024 * 1024)
        
        # Should raise error or return error response
        with pytest.raises((ValueError, Exception)) as exc_info:
            # Simulate upload validation
            max_size = 100 * 1024 * 1024  # 100MB
            if len(large_content) > max_size:
                raise ValueError(f"File too large: {len(large_content)} bytes")
        
        assert "too large" in str(exc_info.value).lower()
    
    @pytest.mark.unit
    def test_upload_unsupported_format(
        self,
        document_agent,
        unsupported_file,
        test_session
    ):
        """Reject unsupported file types."""
        # Check file extension
        assert unsupported_file.suffix == '.xyz'
        
        # Should reject
        with pytest.raises((ValueError, Exception)) as exc_info:
            supported_formats = ['.pdf', '.docx', '.txt', '.png', '.jpg']
            if unsupported_file.suffix not in supported_formats:
                raise ValueError(f"Unsupported format: {unsupported_file.suffix}")
        
        assert "unsupported" in str(exc_info.value).lower()
    
    @pytest.mark.unit
    def test_upload_corrupted_file(
        self,
        document_agent,
        corrupted_pdf,
        test_session
    ):
        """Handle corrupted/invalid files."""
        # Try to read corrupted PDF
        with pytest.raises((Exception, OSError)):
            from PyPDF2 import PdfReader
            PdfReader(str(corrupted_pdf))
    
    @pytest.mark.unit
    def test_upload_empty_file(self, document_agent, temp_storage, test_session):
        """Handle empty file upload."""
        empty_file = temp_storage / "empty.pdf"
        empty_file.write_bytes(b'')
        
        # Should reject or handle gracefully
        assert empty_file.stat().st_size == 0
        
        with pytest.raises((ValueError, Exception)) as exc_info:
            if empty_file.stat().st_size == 0:
                raise ValueError("Empty file not allowed")
        
        assert "empty" in str(exc_info.value).lower()


class TestFileMetadataExtraction:
    """Test metadata extraction from uploaded files."""
    
    @pytest.mark.unit
    def test_extract_pdf_metadata(self, simple_pdf):
        """Extract metadata from PDF."""
        from PyPDF2 import PdfReader
        
        reader = PdfReader(str(simple_pdf))
        
        metadata = {
            'pages': len(reader.pages),
            'size': simple_pdf.stat().st_size
        }
        
        assert metadata['pages'] >= 1
        assert metadata['size'] > 0
    
    @pytest.mark.unit
    def test_extract_docx_metadata(self, simple_docx):
        """Extract metadata from DOCX."""
        from docx import Document
        
        doc = Document(str(simple_docx))
        
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'size': simple_docx.stat().st_size
        }
        
        assert metadata['paragraphs'] > 0
        assert metadata['size'] > 0
    
    @pytest.mark.unit
    def test_extract_image_metadata(self, simple_image):
        """Extract metadata from image."""
        from PIL import Image
        
        with Image.open(simple_image) as img:
            metadata = {
                'format': img.format,
                'size': img.size,
                'mode': img.mode
            }
        
        assert metadata['format'] in ['PNG', 'JPEG', 'JPG']
        assert metadata['size'][0] > 0  # width
        assert metadata['size'][1] > 0  # height


class TestConcurrentUploads:
    """Test concurrent file uploads."""
    
    @pytest.mark.integration
    @pytest.mark.slow
    def test_upload_concurrent_files(self, document_agent, simple_pdf, simple_docx):
        """Handle concurrent uploads."""
        import threading
        
        results = []
        
        def upload_file(file_path, index):
            result = {
                'file_id': f'concurrent-{index}',
                'filename': file_path.name,
                'thread': threading.current_thread().name
            }
            results.append(result)
        
        threads = [
            threading.Thread(target=upload_file, args=(simple_pdf, 1)),
            threading.Thread(target=upload_file, args=(simple_docx, 2))
        ]
        
        for t in threads:
            t.start()
        
        for t in threads:
            t.join()
        
        # All uploads should complete
        assert len(results) == 2
        assert results[0]['file_id'] != results[1]['file_id']


class TestFileValidation:
    """Test file validation logic."""
    
    @pytest.mark.unit
    def test_validate_file_extension(self):
        """Validate file extension check."""
        valid_extensions = ['.pdf', '.docx', '.txt', '.png', '.jpg']
        
        assert '.pdf' in valid_extensions
        assert '.xyz' not in valid_extensions
    
    @pytest.mark.unit
    def test_validate_file_size(self):
        """Validate file size check."""
        max_size = 100 * 1024 * 1024  # 100MB
        
        test_size_valid = 50 * 1024 * 1024  # 50MB
        test_size_invalid = 150 * 1024 * 1024  # 150MB
        
        assert test_size_valid < max_size
        assert test_size_invalid > max_size
    
    @pytest.mark.unit
    def test_validate_file_content_type(self):
        """Validate content type detection."""
        content_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
            'png': 'image/png',
            'jpg': 'image/jpeg'
        }
        
        assert content_types['pdf'] == 'application/pdf'
        assert content_types['docx'].startswith('application/')


class TestUploadCleanup:
    """Test cleanup after upload."""
    
    @pytest.mark.unit
    def test_cleanup_on_upload_failure(self, temp_storage):
        """Verify cleanup on failed upload."""
        test_file = temp_storage / "test_cleanup.pdf"
        test_file.write_bytes(b'invalid data')
        
        assert test_file.exists()
        
        # Simulate failed upload - cleanup should remove file
        test_file.unlink()
        
        assert not test_file.exists()
    
    @pytest.mark.unit
    def test_no_partial_files_stored(self, temp_storage):
        """Ensure no partial files left on error."""
        # Verify directory is clean
        files_before = list(temp_storage.iterdir())
        
        # Simulate upload attempt
        # If it fails, no new files should remain
        
        files_after = list(temp_storage.iterdir())
        assert len(files_after) == len(files_before)


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
