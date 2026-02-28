"""
Generate synthetic test data files for agent tests.
Run once: python generate_test_data.py
"""
import csv
from pathlib import Path

TEST_DATA_DIR = Path(__file__).parent

def make_csv():
    path = TEST_DATA_DIR / "sales_data.csv"
    rows = [
        ["Date", "Product", "Region", "Revenue", "Units"],
        ["2024-01-05", "Widget A", "North", 1500.00, 30],
        ["2024-01-08", "Widget B", "South", 800.00, 16],
        ["2024-01-12", "Widget C", "East", 2200.00, 44],
        ["2024-01-15", "Widget A", "West", 950.00, 19],
        ["2024-01-18", "Widget B", "North", 430.00, 8],
        ["2024-01-22", "Widget C", "South", 3100.00, 62],
        ["2024-01-25", "Widget A", "East", 1750.00, 35],
        ["2024-01-28", "Widget B", "West", 600.00, 12],
        ["2024-02-02", "Widget C", "North", 2800.00, 56],
        ["2024-02-05", "Widget A", "South", 1200.00, 24],
        ["2024-02-09", "Widget B", "East", 550.00, 11],
        ["2024-02-12", "Widget C", "West", 3400.00, 68],
        ["2024-02-16", "Widget A", "North", 1650.00, 33],
        ["2024-02-19", "Widget B", "South", 720.00, 14],
        ["2024-02-23", "Widget C", "East", 2950.00, 59],
        ["2024-03-01", "Widget A", "West", 1100.00, 22],
        ["2024-03-04", "Widget B", "North", 480.00, 9],
        ["2024-03-08", "Widget C", "South", 3250.00, 65],
        ["2024-03-12", "Widget A", "East", 1800.00, 36],
        ["2024-03-15", "Widget B", "West", 670.00, 13],
        ["2024-03-19", "Widget C", "North", 2650.00, 53],
        ["2024-03-22", "Widget A", "South", 1350.00, 27],
        ["2024-03-26", "Widget B", "East", 510.00, 10],
        ["2024-03-29", "Widget C", "West", 3700.00, 74],
        ["2024-04-02", "Widget A", "North", 1900.00, 38],
        ["2024-04-05", "Widget B", "South", 760.00, 15],
        ["2024-04-09", "Widget C", "East", 2500.00, 50],
        ["2024-04-12", "Widget A", "West", 1050.00, 21],
        ["2024-04-16", "Widget B", "North", 520.00, 10],
        ["2024-04-19", "Widget C", "South", 3150.00, 63],
        ["2024-04-23", "Widget A", "East", 1700.00, 34],
        ["2024-04-26", "Widget B", "West", 640.00, 12],
        ["2024-05-01", "Widget C", "North", 2900.00, 58],
        ["2024-05-04", "Widget A", "South", 1250.00, 25],
        ["2024-05-08", "Widget B", "East", 580.00, 11],
        ["2024-05-11", "Widget C", "West", 3600.00, 72],
        ["2024-05-15", "Widget A", "North", 1850.00, 37],
        ["2024-05-18", "Widget B", "South", 700.00, 14],
        ["2024-05-22", "Widget C", "East", 2750.00, 55],
        ["2024-05-25", "Widget A", "West", 1150.00, 23],
        ["2024-05-29", "Widget B", "North", 490.00, 9],
        ["2024-06-01", "Widget C", "South", 3050.00, 61],
        ["2024-06-05", "Widget A", "East", 1950.00, 39],
        ["2024-06-08", "Widget B", "West", 620.00, 12],
        ["2024-06-12", "Widget C", "North", 2600.00, 52],
        ["2024-06-15", "Widget A", "South", 1400.00, 28],
        ["2024-06-19", None,        "East", 530.00,  10],   # missing product
        ["2024-06-22", "Widget C", "West", 3800.00, 76],
        ["2024-06-26", "Widget A", "North", 1600.00, 32],
        ["2024-06-29", "Widget B", "South", 740.00, 14],
    ]
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f"✅ Created {path}")


def make_xlsx():
    try:
        import openpyxl
    except ImportError:
        print("⚠️  openpyxl not installed, skipping XLSX generation")
        return

    path = TEST_DATA_DIR / "employees.xlsx"
    wb = openpyxl.Workbook()

    # Sheet 1: Employees
    ws1 = wb.active
    ws1.title = "Employees"
    ws1.append(["EmployeeID", "Name", "DepartmentID", "Salary", "Hire Date", "Status"])
    employees = [
        (101, "Alice Johnson",    1, 85000, "2021-03-15", "Active"),
        (102, "Bob Smith",        2, 72000, "2020-07-01", "Active"),
        (103, "Carol White",      1, 91000, "2019-11-20", "Active"),
        (104, "David Brown",      3, 68000, "2022-01-10", "Active"),
        (105, "Eve Davis",        2, 79000, "2021-08-05", "Inactive"),
        (106, "Frank Martinez",   None,    64000, "2023-03-01", "Active"),
        (107, "Grace Lee",        1, 95000, "2018-06-15", "Active"),
        (108, "Hank Wilson",      3, 71000, "2022-09-20", "Active"),
        (109, "Iris Taylor",      2, 83000, "2020-02-14", "Active"),
        (110, "Jack Anderson",    None,    77000, "2021-05-30", "Inactive"),
    ]
    for emp in employees:
        ws1.append(list(emp))

    # Sheet 2: Departments
    ws2 = wb.create_sheet("Departments")
    ws2.append(["DepartmentID", "DepartmentName", "Manager", "Budget"])
    departments = [
        (1, "Engineering",  "Carol White",   500000),
        (2, "Marketing",    "Iris Taylor",   300000),
        (3, "Operations",   "Hank Wilson",   250000),
    ]
    for dept in departments:
        ws2.append(list(dept))

    wb.save(path)
    print(f"✅ Created {path}")


def make_docx():
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("⚠️  python-docx not installed, skipping DOCX generation")
        return

    path = TEST_DATA_DIR / "sample_report.docx"
    doc = Document()

    doc.add_heading("Q2 2024 Business Performance Report", 0)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Q2 2024 saw strong revenue growth across all product lines, with total revenue "
        "reaching $47.5 million — a 12% increase year-over-year. Widget C continued to "
        "be our top performer, accounting for 58% of total revenue. The North and South "
        "regions showed the strongest growth, while the West region recovered from Q1 dip."
    )

    doc.add_heading("Revenue by Product", level=1)
    doc.add_paragraph(
        "Widget A contributed $8.2 million, Widget B contributed $4.1 million, and "
        "Widget C led with $35.2 million. Combined units sold reached 1,240 across all "
        "product lines, with an average selling price of $38.30."
    )

    doc.add_heading("Regional Breakdown", level=1)
    doc.add_paragraph(
        "The North region generated $14.3 million (30%), the South region $12.8 million (27%), "
        "the East region $11.9 million (25%), and the West region $8.5 million (18%)."
    )

    doc.add_heading("Financial Table", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Product"
    hdr[1].text = "Revenue ($M)"
    hdr[2].text = "Units"
    hdr[3].text = "Avg Price ($)"
    data = [
        ("Widget A", "8.2", "430", "19.07"),
        ("Widget B", "4.1", "215", "19.07"),
        ("Widget C", "35.2", "595", "59.16"),
        ("Total",    "47.5", "1,240", "38.31"),
    ]
    for d in data:
        row = table.add_row().cells
        for i, val in enumerate(d):
            row[i].text = val

    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph(
        "1. Increase Widget C production capacity to meet growing demand.\n"
        "2. Launch targeted campaign in the West region to boost market share.\n"
        "3. Investigate Widget B underperformance and explore pricing adjustments."
    )

    doc.save(path)
    print(f"✅ Created {path}")


def make_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as pdfcanvas
    except ImportError:
        print("⚠️  reportlab not installed, skipping PDF generation")
        return

    path = TEST_DATA_DIR / "sample_invoice.pdf"
    c = pdfcanvas.Canvas(str(path), pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 20)
    c.drawString(50, height - 60, "INVOICE")

    c.setFont("Helvetica", 11)
    c.drawString(50, height - 100, "Invoice No: INV-2024-00123")
    c.drawString(50, height - 118, "Date: June 30, 2024")
    c.drawString(50, height - 136, "Due Date: July 30, 2024")

    c.drawString(50, height - 175, "Bill To:")
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, height - 193, "Acme Corporation")
    c.setFont("Helvetica", 11)
    c.drawString(50, height - 211, "123 Business Ave, New York, NY 10001")

    c.drawString(50, height - 250, "From:")
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, height - 268, "Orbimesh Solutions Pvt. Ltd.")
    c.setFont("Helvetica", 11)
    c.drawString(50, height - 286, "456 Tech Park, Hyderabad, TS 500081")

    # Table header
    y = height - 340
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, "Description")
    c.drawString(320, y, "Qty")
    c.drawString(380, y, "Unit Price")
    c.drawString(470, y, "Total")
    c.line(50, y - 5, 550, y - 5)

    c.setFont("Helvetica", 11)
    items = [
        ("AI Orchestration Platform — Monthly License", "1", "$4,500.00", "$4,500.00"),
        ("Custom Agent Development (hours)", "20",  "$150.00",   "$3,000.00"),
        ("Support & Maintenance",           "1",   "$500.00",   "$500.00"),
    ]
    for i, (desc, qty, unit, total) in enumerate(items):
        row_y = y - 25 - i * 22
        c.drawString(50, row_y, desc)
        c.drawString(320, row_y, qty)
        c.drawString(380, row_y, unit)
        c.drawString(470, row_y, total)

    # Totals
    tot_y = y - 25 - len(items) * 22 - 20
    c.line(50, tot_y + 10, 550, tot_y + 10)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(380, tot_y - 5,  "Subtotal:")
    c.drawString(470, tot_y - 5,  "$8,000.00")
    c.drawString(380, tot_y - 25, "Tax (18%):")
    c.drawString(470, tot_y - 25, "$1,440.00")
    c.drawString(380, tot_y - 50, "Total Due:")
    c.drawString(470, tot_y - 50, "$9,440.00")

    c.setFont("Helvetica", 10)
    c.drawString(50, 80, "Payment Terms: Net 30. Bank Transfer or UPI accepted.")
    c.drawString(50, 65, "Thank you for your business!")

    c.save()
    print(f"✅ Created {path}")


if __name__ == "__main__":
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    make_csv()
    make_xlsx()
    make_docx()
    make_pdf()
    print("\nAll test data generated successfully.")
