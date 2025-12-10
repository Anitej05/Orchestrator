"""
Intelligent Document Editor - Handles complex document modifications.

This editor goes beyond simple add operations to intelligently modify existing content.
It analyzes the document state and applies minimal changes to achieve the desired result.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from typing import Dict, List, Any, Optional, Tuple
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class IntelligentDocumentEditor:
    """
    Intelligent document editor that understands document state and applies smart modifications.
    """
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
    
    def analyze_current_state(self) -> Dict[str, Any]:
        """
        Analyze the current state of the document.
        Returns detailed information about content, structure, and formatting.
        """
        state = {
            'total_paragraphs': len(self.doc.paragraphs),
            'paragraphs': [],
            'lists': [],
            'headings': [],
            'tables': len(self.doc.tables),
            'styles_used': {}
        }
        
        current_list = None
        list_index = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            para_info = {
                'index': i,
                'text': para.text,
                'style': para.style.name,
                'is_list_item': self._is_list_item(para),
                'list_type': self._get_list_type(para),
                'alignment': str(para.alignment) if para.alignment else None,
                'runs': []
            }
            
            # Analyze runs for formatting
            for run in para.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_color': self._get_color_name(run.font.color) if run.font.color else None
                }
                para_info['runs'].append(run_info)
            
            state['paragraphs'].append(para_info)
            
            # Track lists
            if para_info['is_list_item']:
                if current_list is None or current_list['type'] != para_info['list_type']:
                    current_list = {
                        'index': list_index,
                        'type': para_info['list_type'],
                        'start_para': i,
                        'items': []
                    }
                    state['lists'].append(current_list)
                    list_index += 1
                current_list['items'].append({
                    'para_index': i,
                    'text': para.text
                })
            else:
                current_list = None
            
            # Track headings
            if 'Heading' in para.style.name:
                state['headings'].append({
                    'level': para.style.name,
                    'text': para.text,
                    'para_index': i
                })
            
            # Track styles
            style_name = para.style.name
            state['styles_used'][style_name] = state['styles_used'].get(style_name, 0) + 1
        
        return state
    
    def _is_list_item(self, paragraph) -> bool:
        """Check if a paragraph is a list item."""
        try:
            # Check for numbering XML
            if paragraph._element.pPr is not None:
                numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    return True
            
            # Also check if style name indicates a list
            style_name = paragraph.style.name.lower()
            if 'list' in style_name or 'bullet' in style_name or 'number' in style_name:
                return True
            
            return False
        except:
            return False
    
    def _get_list_type(self, paragraph) -> Optional[str]:
        """Determine if list is numbered or bulleted."""
        try:
            # Check style name first
            style_name = paragraph.style.name.lower()
            if 'number' in style_name:
                return 'numbered'
            elif 'bullet' in style_name:
                return 'bulleted'
            
            # Check XML numbering
            if paragraph._element.pPr is not None:
                numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    numId = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                    if numId is not None:
                        # Check the text to determine type
                        text = paragraph.text.strip()
                        if text and text[0].isdigit():
                            return 'numbered'
                        else:
                            return 'bulleted'
            
            # Check text content as fallback
            text = paragraph.text.strip()
            if text:
                if text.startswith('-') or text.startswith('•') or text.startswith('·'):
                    return 'bulleted'
                elif text[0].isdigit() and (len(text) > 1 and text[1] in '.):'):
                    return 'numbered'
            
            return None
        except:
            return None
    
    def _get_color_name(self, color) -> Optional[str]:
        """Convert RGB color to common color name."""
        if not color or not color.rgb:
            return None
        
        rgb = color.rgb
        # Simple color mapping
        color_map = {
            'FF0000': 'red',
            '00FF00': 'green',
            '0000FF': 'blue',
            'FFFF00': 'yellow',
            'FFA500': 'orange',
            '800080': 'purple',
            '000000': 'black',
            'FFFFFF': 'white'
        }
        return color_map.get(str(rgb).upper(), f'#{rgb}')
    
    def convert_list_style(self, from_type: str, to_type: str, list_indices: Optional[List[int]] = None):
        """
        Convert list style from one type to another.
        
        Args:
            from_type: 'bulleted' or 'numbered'
            to_type: 'bulleted' or 'numbered'
            list_indices: Specific list indices to convert (None = all lists)
        """
        state = self.analyze_current_state()
        lists_to_convert = state['lists']
        
        if list_indices is not None:
            lists_to_convert = [l for l in lists_to_convert if l['index'] in list_indices]
        
        for list_info in lists_to_convert:
            if list_info['type'] == from_type:
                for item in list_info['items']:
                    para = self.doc.paragraphs[item['para_index']]
                    self._set_list_style(para, to_type)
        
        logger.info(f"Converted {len(lists_to_convert)} lists from {from_type} to {to_type}")
    
    def _set_list_style(self, paragraph, list_type: str):
        """Set the list style for a paragraph."""
        # Remove existing numbering
        if paragraph._element.pPr is not None:
            numPr = paragraph._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                paragraph._element.pPr.remove(numPr)
        
        # Add new numbering
        if list_type == 'numbered':
            # Add numbered list formatting
            self._apply_numbering(paragraph, num_id=1, ilvl=0)  # num_id=1 is typically numbered
        elif list_type == 'bulleted':
            # Add bulleted list formatting
            self._apply_numbering(paragraph, num_id=2, ilvl=0)  # num_id=2 is typically bulleted
    
    def _apply_numbering(self, paragraph, num_id: int, ilvl: int):
        """Apply numbering to a paragraph."""
        pPr = paragraph._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        
        ilvl_element = OxmlElement('w:ilvl')
        ilvl_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(ilvl))
        numPr.append(ilvl_element)
        
        numId_element = OxmlElement('w:numId')
        numId_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(num_id))
        numPr.append(numId_element)
        
        pPr.append(numPr)
    
    def remove_paragraphs_by_pattern(self, pattern: str = None, indices: List[int] = None):
        """
        Remove paragraphs matching a pattern or at specific indices.
        
        Args:
            pattern: Text pattern to match (removes paragraphs containing this text)
            indices: Specific paragraph indices to remove
        """
        if indices:
            # Remove by index (in reverse order to maintain indices)
            for i in sorted(indices, reverse=True):
                if 0 <= i < len(self.doc.paragraphs):
                    p = self.doc.paragraphs[i]
                    p._element.getparent().remove(p._element)
                    logger.info(f"Removed paragraph at index {i}")
        
        elif pattern:
            # Remove by pattern
            removed_count = 0
            for para in list(self.doc.paragraphs):  # Create a copy to iterate
                if pattern.lower() in para.text.lower():
                    para._element.getparent().remove(para._element)
                    removed_count += 1
            logger.info(f"Removed {removed_count} paragraphs matching pattern '{pattern}'")
    
    def modify_text_formatting(
        self,
        target_text: Optional[str] = None,
        target_style: Optional[str] = None,
        para_indices: Optional[List[int]] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None,
        color: Optional[str] = None,
        font_size: Optional[int] = None
    ):
        """
        Modify formatting of existing text.
        
        Args:
            target_text: Text to find and modify
            target_style: Style name to target (e.g., 'Heading 2')
            para_indices: Specific paragraph indices to modify
            bold, italic, underline, color, font_size: Formatting to apply
        """
        paragraphs_to_modify = []
        
        if para_indices:
            paragraphs_to_modify = [self.doc.paragraphs[i] for i in para_indices if i < len(self.doc.paragraphs)]
        elif target_style:
            paragraphs_to_modify = [p for p in self.doc.paragraphs if p.style.name == target_style]
        elif target_text:
            paragraphs_to_modify = [p for p in self.doc.paragraphs if target_text.lower() in p.text.lower()]
        
        modified_count = 0
        for para in paragraphs_to_modify:
            # If paragraph has no runs but has text, we need to handle it differently
            if not para.runs and para.text:
                # Create a run with the text
                para.add_run(para.text)
                # Clear the original text
                para.text = ''
            
            # Now modify all runs
            if para.runs:
                for run in para.runs:
                    if bold is not None:
                        run.bold = bold
                    if italic is not None:
                        run.italic = italic
                    if underline is not None:
                        run.underline = underline
                    if color:
                        run.font.color.rgb = self._parse_color(color)
                    if font_size:
                        run.font.size = Pt(font_size)
                modified_count += 1
        
        logger.info(f"Modified formatting for {modified_count} paragraphs")
    
    def _parse_color(self, color_input: str) -> RGBColor:
        """
        Parse color input to RGBColor.
        Accepts: color names, hex codes (#RRGGBB), or RGB tuples (r,g,b)
        """
        import re
        
        color_input = color_input.strip()
        
        # Try hex format first (#RRGGBB or RRGGBB)
        hex_match = re.match(r'^#?([0-9A-Fa-f]{6})$', color_input)
        if hex_match:
            hex_color = hex_match.group(1)
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        
        # Try RGB tuple format (r,g,b)
        rgb_match = re.match(r'^\(?(\d+),\s*(\d+),\s*(\d+)\)?$', color_input)
        if rgb_match:
            r, g, b = map(int, rgb_match.groups())
            return RGBColor(r, g, b)
        
        # Common color names (minimal set for fallback)
        basic_colors = {
            'red': (255, 0, 0),
            'green': (0, 255, 0),
            'blue': (0, 0, 255),
            'yellow': (255, 255, 0),
            'orange': (255, 165, 0),
            'purple': (128, 0, 128),
            'brown': (165, 42, 42),
            'pink': (255, 192, 203),
            'gray': (128, 128, 128),
            'grey': (128, 128, 128),
            'black': (0, 0, 0),
            'white': (255, 255, 255),
            'violet': (238, 130, 238),
            'cyan': (0, 255, 255),
            'magenta': (255, 0, 255)
        }
        
        color_lower = color_input.lower()
        if color_lower in basic_colors:
            r, g, b = basic_colors[color_lower]
            return RGBColor(r, g, b)
        
        # If not recognized, log warning and return black
        logger.warning(f"Color '{color_input}' not recognized, using black. Use hex (#RRGGBB) or RGB (r,g,b) for custom colors.")
        return RGBColor(0, 0, 0)
    
    def remove_all_formatting(self, para_indices: Optional[List[int]] = None):
        """Remove all formatting from paragraphs (reset to default)."""
        paragraphs = [self.doc.paragraphs[i] for i in para_indices] if para_indices else self.doc.paragraphs
        
        for para in paragraphs:
            for run in para.runs:
                run.bold = False
                run.italic = False
                run.underline = False
                run.font.color.rgb = None
                run.font.size = None
        
        logger.info(f"Removed formatting from {len(paragraphs)} paragraphs")
    
    def delete_tables(self, table_indices: Optional[List[int]] = None):
        """Delete tables from the document."""
        if table_indices is None:
            # Delete all tables
            for table in list(self.doc.tables):
                table._element.getparent().remove(table._element)
            logger.info(f"Deleted all tables")
        else:
            # Delete specific tables (in reverse order)
            for i in sorted(table_indices, reverse=True):
                if 0 <= i < len(self.doc.tables):
                    table = self.doc.tables[i]
                    table._element.getparent().remove(table._element)
            logger.info(f"Deleted {len(table_indices)} tables")
    
    def modify_table(self, table_index: int, modifications: Dict[str, Any]):
        """
        Modify an existing table.
        
        Args:
            table_index: Index of table to modify
            modifications: Dict with keys like 'add_row', 'delete_row', 'update_cell', etc.
        """
        if table_index >= len(self.doc.tables):
            logger.error(f"Table index {table_index} out of range")
            return
        
        table = self.doc.tables[table_index]
        
        if 'add_row' in modifications:
            row = table.add_row()
            logger.info(f"Added row to table {table_index}")
        
        if 'delete_row' in modifications:
            row_idx = modifications['delete_row']
            if 0 <= row_idx < len(table.rows):
                table._element.remove(table.rows[row_idx]._element)
                logger.info(f"Deleted row {row_idx} from table {table_index}")
        
        if 'update_cell' in modifications:
            cell_info = modifications['update_cell']
            row_idx = cell_info.get('row', 0)
            col_idx = cell_info.get('col', 0)
            text = cell_info.get('text', '')
            if 0 <= row_idx < len(table.rows) and 0 <= col_idx < len(table.columns):
                table.cell(row_idx, col_idx).text = text
                logger.info(f"Updated cell ({row_idx}, {col_idx}) in table {table_index}")
    
    def change_style(self, para_indices: List[int], new_style: str):
        """Change the style of specific paragraphs."""
        for i in para_indices:
            if 0 <= i < len(self.doc.paragraphs):
                try:
                    self.doc.paragraphs[i].style = new_style
                    logger.info(f"Changed style of paragraph {i} to {new_style}")
                except Exception as e:
                    logger.error(f"Failed to change style of paragraph {i}: {e}")
    
    def remove_all_lists(self):
        """Remove all list formatting from the document."""
        for para in self.doc.paragraphs:
            if self._is_list_item(para):
                # Remove numbering
                if para._element.pPr is not None:
                    numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numPr is not None:
                        para._element.pPr.remove(numPr)
        logger.info("Removed all list formatting")
    
    def clear_document_content(self, keep_structure: bool = False):
        """
        Clear all content from document.
        
        Args:
            keep_structure: If True, keeps headings and structure but removes content
        """
        if keep_structure:
            # Remove only body text, keep headings
            for para in list(self.doc.paragraphs):
                if 'Heading' not in para.style.name:
                    para._element.getparent().remove(para._element)
        else:
            # Remove everything
            for para in list(self.doc.paragraphs):
                para._element.getparent().remove(para._element)
            for table in list(self.doc.tables):
                table._element.getparent().remove(table._element)
        
        logger.info(f"Cleared document content (keep_structure={keep_structure})")
    
    def find_and_delete(self, search_text: str, case_sensitive: bool = False):
        """Find and delete all paragraphs containing specific text."""
        deleted_count = 0
        for para in list(self.doc.paragraphs):
            text = para.text if case_sensitive else para.text.lower()
            search = search_text if case_sensitive else search_text.lower()
            
            if search in text:
                para._element.getparent().remove(para._element)
                deleted_count += 1
        
        logger.info(f"Deleted {deleted_count} paragraphs containing '{search_text}'")
        return deleted_count
    
    def replace_all_text(self, old_text: str, new_text: str, case_sensitive: bool = False):
        """Replace all occurrences of text in the document."""
        replaced_count = 0
        for para in self.doc.paragraphs:
            if case_sensitive:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            replaced_count += 1
            else:
                if old_text.lower() in para.text.lower():
                    for run in para.runs:
                        # Case-insensitive replace
                        import re
                        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                        if pattern.search(run.text):
                            run.text = pattern.sub(new_text, run.text)
                            replaced_count += 1
        
        logger.info(f"Replaced {replaced_count} occurrences of '{old_text}' with '{new_text}'")
        return replaced_count
    
    def remove_empty_paragraphs(self):
        """Remove all empty paragraphs from the document."""
        removed_count = 0
        for para in list(self.doc.paragraphs):
            if not para.text.strip():
                para._element.getparent().remove(para._element)
                removed_count += 1
        
        logger.info(f"Removed {removed_count} empty paragraphs")
        return removed_count
    
    def detect_pseudo_lists(self) -> List[Dict[str, Any]]:
        """
        Detect paragraphs that look like lists but don't have Word list formatting.
        Returns list of pseudo-list items with their indices and types.
        """
        import re
        pseudo_lists = []
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Skip if already a real list
            if self._is_list_item(para):
                continue
            
            # Check for common list markers
            marker_type = None
            if text.startswith(('-', '•', '·', '◦', '▪', '▫')):
                marker_type = 'bulleted'
            elif len(text) > 1 and text[0].isdigit() and text[1] in '.):':
                marker_type = 'numbered'
            elif re.match(r'^[a-zA-Z][.):]\s', text):
                marker_type = 'lettered'
            
            if marker_type:
                pseudo_lists.append({
                    'para_index': i,
                    'text': text,
                    'marker_type': marker_type,
                    'style': para.style.name
                })
        
        logger.info(f"Detected {len(pseudo_lists)} pseudo-list items")
        return pseudo_lists
    
    def convert_text_to_list(self, para_indices: Optional[List[int]] = None, list_type: str = 'numbered'):
        """
        Convert plain text paragraphs (with list markers) to actual Word lists.
        
        Args:
            para_indices: Specific paragraph indices to convert (None = auto-detect)
            list_type: 'numbered' or 'bulleted'
        """
        import re
        
        # Auto-detect if not specified
        if para_indices is None:
            pseudo_lists = self.detect_pseudo_lists()
            para_indices = [item['para_index'] for item in pseudo_lists]
        
        converted_count = 0
        for i in para_indices:
            if i >= len(self.doc.paragraphs):
                continue
            
            para = self.doc.paragraphs[i]
            text = para.text.strip()
            
            # Remove list markers from text
            if text.startswith(('-', '•', '·', '◦', '▪', '▫')):
                text = text[1:].strip()
            elif text and text[0].isdigit():
                # Remove number prefix (e.g., "1. " or "1) ")
                text = re.sub(r'^\d+[.):]\s*', '', text)
            elif re.match(r'^[a-zA-Z][.):]\s', text):
                # Remove letter prefix (e.g., "a. " or "A) ")
                text = re.sub(r'^[a-zA-Z][.):]\s*', '', text)
            
            # Clear existing runs and set new text
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            
            # Apply Word list style
            if list_type == 'numbered':
                para.style = 'List Number'
                self._apply_numbering(para, num_id=1, ilvl=0)
            else:
                para.style = 'List Bullet'
                self._apply_numbering(para, num_id=2, ilvl=0)
            
            converted_count += 1
        
        logger.info(f"Converted {converted_count} text paragraphs to {list_type} lists")
        return converted_count
    
    def move_paragraph(self, from_index: int, to_index: int):
        """
        Move a paragraph from one position to another.
        
        Args:
            from_index: Source paragraph index
            to_index: Destination index (paragraph will be inserted before this index)
        """
        if from_index >= len(self.doc.paragraphs) or from_index < 0:
            logger.error(f"Invalid from_index: {from_index}")
            return 0
        
        # Get the paragraph to move
        para_to_move = self.doc.paragraphs[from_index]
        
        # Store paragraph properties
        text = para_to_move.text
        style = para_to_move.style
        
        # Store run formatting
        runs_data = []
        for run in para_to_move.runs:
            runs_data.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            })
        
        # Remove the paragraph
        para_to_move._element.getparent().remove(para_to_move._element)
        
        # Insert at new location
        if to_index >= len(self.doc.paragraphs):
            # Add at end
            new_para = self.doc.add_paragraph()
        else:
            # Insert before to_index
            target_para = self.doc.paragraphs[to_index]
            new_para = target_para.insert_paragraph_before()
        
        # Apply style
        new_para.style = style
        
        # Recreate runs with formatting
        for run_data in runs_data:
            run = new_para.add_run(run_data['text'])
            run.bold = run_data['bold']
            run.italic = run_data['italic']
            run.underline = run_data['underline']
            if run_data['font_size']:
                run.font.size = run_data['font_size']
            if run_data['font_color']:
                run.font.color.rgb = run_data['font_color']
        
        logger.info(f"Moved paragraph from index {from_index} to {to_index}")
        return 1
    
    def insert_paragraph_at(self, index: int, text: str, style: str = 'Normal', **formatting):
        """
        Insert a paragraph at a specific index.
        
        Args:
            index: Position to insert (inserts before this index)
            text: Paragraph text
            style: Paragraph style
            **formatting: bold, italic, underline, color, font_size
        """
        if index >= len(self.doc.paragraphs):
            # Add at end
            new_para = self.doc.add_paragraph(text, style=style)
        else:
            # Insert before index
            target_para = self.doc.paragraphs[index]
            new_para = target_para.insert_paragraph_before(text, style=style)
        
        # Apply formatting
        for run in new_para.runs:
            if formatting.get('bold'):
                run.bold = True
            if formatting.get('italic'):
                run.italic = True
            if formatting.get('underline'):
                run.underline = True
            if formatting.get('color'):
                run.font.color.rgb = self._parse_color(formatting['color'])
            if formatting.get('font_size'):
                run.font.size = Pt(formatting['font_size'])
        
        logger.info(f"Inserted paragraph at index {index}")
        return 1
    
    def save(self, output_path: Optional[str] = None):
        """Save the document."""
        save_path = output_path or self.file_path
        self.doc.save(save_path)
        logger.info(f"Document saved to {save_path}")
        return save_path
