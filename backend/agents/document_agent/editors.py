"""
Document Agent - Editors Module

Combines AdvancedDocumentEditor and IntelligentDocumentEditor for efficient document editing.
Optimized for cloud deployment with minimal memory footprint.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from typing import Dict, List, Any, Optional, Tuple
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentEditor:
    """
    Unified document editor combining advanced and intelligent editing capabilities.
    Optimized for cloud deployment with lazy loading and memory efficiency.
    """

    # Color mapping
    COLORS = {
        'red': (255, 0, 0), 'blue': (0, 0, 255), 'green': (0, 128, 0),
        'yellow': (255, 255, 0), 'orange': (255, 165, 0), 'purple': (128, 0, 128),
        'black': (0, 0, 0), 'white': (255, 255, 255), 'gray': (128, 128, 128),
        'grey': (128, 128, 128), 'cyan': (0, 255, 255), 'magenta': (255, 0, 255),
        'brown': (165, 42, 42), 'pink': (255, 192, 203), 'navy': (0, 0, 128),
    }

    # Style mapping
    STYLE_MAPPING = {
        'subheading': 'Heading 2', 'subtitle': 'Heading 2', 'heading': 'Heading 1',
        'title': 'Title', 'h1': 'Heading 1', 'h2': 'Heading 2', 'h3': 'Heading 3',
        'normal': 'Normal', 'body': 'Normal'
    }

    def __init__(self, file_path: str):
        """Initialize editor with file validation."""
        if not file_path or not isinstance(file_path, str) or file_path.strip() == '':
            raise ValueError(f"Invalid file_path: {repr(file_path)}")

        self.file_path = file_path

        if not Path(file_path).exists():
            raise FileNotFoundError(f"Document not found at: {file_path}")

        try:
            self.doc = Document(file_path)
            logger.info(f"Loaded document: {file_path}")
        except Exception as e:
            raise RuntimeError(f"Failed to open document {file_path}: {str(e)}") from e

    def save(self):
        """Save document to disk."""
        try:
            self.doc.save(self.file_path)
            logger.info(f"Document saved: {self.file_path}")
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            raise

    def analyze_state(self) -> Dict[str, Any]:
        """Analyze current document state (intelligence)."""
        state = {
            'total_paragraphs': len(self.doc.paragraphs),
            'total_tables': len(self.doc.tables),
            'paragraphs': [],
            'headings': [],
            'styles_used': {}
        }

        for i, para in enumerate(self.doc.paragraphs):
            para_info = {
                'index': i,
                'text': para.text[:100],  # Limit to 100 chars
                'style': para.style.name,
                'is_heading': 'Heading' in para.style.name
            }
            state['paragraphs'].append(para_info)

            if 'Heading' in para.style.name:
                state['headings'].append({
                    'level': para.style.name,
                    'text': para.text
                })

            style_name = para.style.name
            state['styles_used'][style_name] = state['styles_used'].get(style_name, 0) + 1

        return state

    # ========== TEXT FORMATTING ==========

    def format_text(self, text: str, **kwargs) -> str:
        """Apply formatting to text (advanced)."""
        for para in self.doc.paragraphs:
            if text in para.text:
                for run in para.runs:
                    if text in run.text:
                        if kwargs.get('bold'):
                            run.bold = True
                        if kwargs.get('italic'):
                            run.italic = True
                        if kwargs.get('underline'):
                            run.underline = True
                        if 'font_size' in kwargs:
                            run.font.size = Pt(kwargs['font_size'])
                        if 'color' in kwargs:
                            color = self._parse_color(kwargs['color'])
                            run.font.color.rgb = RGBColor(*color)
                return "✓ Text formatted successfully"

        return "✗ Text not found in document"

    def add_heading(self, text: str, level: int = 1) -> str:
        """Add heading to document."""
        try:
            self.doc.add_heading(text, level=level)
            logger.info(f"Added heading: {text}")
            return f"✓ Added heading: {text}"
        except Exception as e:
            logger.error(f"Failed to add heading: {e}")
            return f"✗ Failed to add heading: {e}"

    def add_paragraph(self, text: str, style: str = 'Normal') -> str:
        """Add paragraph to document."""
        try:
            mapped_style = self._map_style(style)
            self.doc.add_paragraph(text, style=mapped_style)
            logger.info(f"Added paragraph with style {mapped_style}")
            return f"✓ Added paragraph"
        except Exception as e:
            logger.error(f"Failed to add paragraph: {e}")
            return f"✗ Failed to add paragraph: {e}"

    def add_table(self, rows: int, cols: int) -> str:
        """Add table to document."""
        try:
            self.doc.add_table(rows=rows, cols=cols)
            logger.info(f"Added table: {rows}x{cols}")
            return f"✓ Added {rows}x{cols} table"
        except Exception as e:
            logger.error(f"Failed to add table: {e}")
            return f"✗ Failed to add table: {e}"

    def replace_text(self, old_text: str, new_text: str) -> str:
        """Replace text in document (intelligence)."""
        count = 0
        for para in self.doc.paragraphs:
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
                count += 1

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
                        count += 1

        if count > 0:
            logger.info(f"Replaced text in {count} locations")
            return f"✓ Replaced text in {count} locations"
        return "✗ Text not found"

    def delete_paragraph(self, index: int) -> str:
        """Delete paragraph by index (intelligence)."""
        try:
            if 0 <= index < len(self.doc.paragraphs):
                p = self.doc.paragraphs[index]._element
                p.getparent().remove(p)
                logger.info(f"Deleted paragraph at index {index}")
                return f"✓ Deleted paragraph"
            return "✗ Invalid paragraph index"
        except Exception as e:
            logger.error(f"Failed to delete paragraph: {e}")
            return f"✗ Failed to delete paragraph: {e}"

    def set_paragraph_style(self, index: int, style: str) -> str:
        """Set style of paragraph (intelligence)."""
        try:
            if 0 <= index < len(self.doc.paragraphs):
                mapped_style = self._map_style(style)
                self.doc.paragraphs[index].style = mapped_style
                logger.info(f"Set paragraph {index} style to {mapped_style}")
                return f"✓ Applied style: {mapped_style}"
            return "✗ Invalid paragraph index"
        except Exception as e:
            logger.error(f"Failed to set style: {e}")
            return f"✗ Failed to set style: {e}"

    def add_image(self, image_path: str, width: Optional[int] = None) -> str:
        """Add image to document."""
        try:
            if not Path(image_path).exists():
                return f"✗ Image not found: {image_path}"

            if width:
                self.doc.add_picture(image_path, width=Inches(width))
            else:
                self.doc.add_picture(image_path)

            logger.info(f"Added image: {image_path}")
            return f"✓ Added image"
        except Exception as e:
            logger.error(f"Failed to add image: {e}")
            return f"✗ Failed to add image: {e}"

    # ========== HELPER METHODS ==========

    def _map_style(self, style: str) -> str:
        """Map style name to Word style."""
        return self.STYLE_MAPPING.get(style.lower(), style)

    def _parse_color(self, color: Any) -> Tuple[int, int, int]:
        """Parse color from string or RGB tuple."""
        if isinstance(color, str):
            return self.COLORS.get(color.lower(), (0, 0, 0))
        elif isinstance(color, (list, tuple)) and len(color) == 3:
            return tuple(color)
        return (0, 0, 0)

    def get_summary(self) -> Dict[str, Any]:
        """Get brief document summary for cloud operations."""
        return {
            'file': Path(self.file_path).name,
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'size_kb': Path(self.file_path).stat().st_size / 1024
        }

    # ========== INTELLIGENT EDITING METHODS ==========

    def analyze_current_state(self) -> Dict[str, Any]:
        """
        Analyze the current state of the document with detailed information.
        Returns information about content, structure, formatting, lists, etc.
        """
        state = {
            'total_paragraphs': len(self.doc.paragraphs),
            'paragraphs': [],
            'lists': [],
            'headings': [],
            'tables': len(self.doc.tables),
            'styles_used': {}
        }
        
        current_list = None
        list_index = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            para_info = {
                'index': i,
                'text': para.text[:200] if para.text else '',
                'style': para.style.name,
                'is_heading': 'Heading' in para.style.name
            }
            
            # Check formatting
            if para.runs:
                run = para.runs[0]
                para_info['bold'] = run.bold if run.bold is not None else False
                para_info['italic'] = run.italic if run.italic is not None else False
                para_info['underline'] = run.underline if run.underline is not None else False
                if run.font.color and run.font.color.rgb:
                    para_info['color'] = self._get_color_name(run.font.color)
            
            # Check if it's a list item
            is_list = self._is_list_item(para)
            para_info['is_list_item'] = is_list
            
            if is_list:
                list_type = self._get_list_type(para)
                if current_list is None or current_list['type'] != list_type:
                    # Start new list
                    current_list = {
                        'index': list_index,
                        'type': list_type,
                        'start_para': i,
                        'items': []
                    }
                    state['lists'].append(current_list)
                    list_index += 1
                
                current_list['items'].append({
                    'para_index': i,
                    'text': para.text[:100]
                })
                current_list['end_para'] = i
            else:
                current_list = None
            
            state['paragraphs'].append(para_info)
            
            # Track headings
            if 'Heading' in para.style.name:
                state['headings'].append({
                    'level': para.style.name,
                    'text': para.text,
                    'index': i
                })
            
            # Track styles
            style_name = para.style.name
            state['styles_used'][style_name] = state['styles_used'].get(style_name, 0) + 1
        
        return state

    def _is_list_item(self, paragraph) -> bool:
        """Check if a paragraph is a list item."""
        try:
            pPr = paragraph._element.pPr
            if pPr is None:
                return False
            numPr = pPr.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            return numPr is not None
        except:
            return False

    def _get_list_type(self, paragraph) -> Optional[str]:
        """Determine if list is numbered or bulleted."""
        try:
            pPr = paragraph._element.pPr
            if pPr is None:
                return None
            
            numPr = pPr.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if numPr is None:
                return None
            
            # Check numbering format from document's numbering.xml
            # For simplicity, check text content
            text = paragraph.text.strip()
            if text:
                first_char = text[0]
                if first_char.isdigit():
                    return 'numbered'
                elif first_char in ['•', '\u2022', '\u25cf', '-', '*']:
                    return 'bulleted'
            
            return 'numbered'  # Default
        except:
            return None

    def _get_color_name(self, color) -> Optional[str]:
        """Convert RGB color to common color name."""
        if not color or not color.rgb:
            return None
        
        rgb = color.rgb
        # Convert to tuple of ints
        r, g, b = rgb[0], rgb[1], rgb[2]
        
        # Match to known colors
        for name, (cr, cg, cb) in self.COLORS.items():
            if abs(r - cr) < 20 and abs(g - cg) < 20 and abs(b - cb) < 20:
                return name
        
        return f'rgb({r},{g},{b})'

    def convert_list_style(self, from_type: str, to_type: str, list_indices: Optional[List[int]] = None):
        """Convert list style from one type to another."""
        try:
            for i, para in enumerate(self.doc.paragraphs):
                if not self._is_list_item(para):
                    continue
                
                current_type = self._get_list_type(para)
                if current_type != from_type:
                    continue
                
                if list_indices and i not in list_indices:
                    continue
                
                # Apply new list style
                self._set_list_style(para, to_type)
            
            logger.info(f\"Converted lists from {from_type} to {to_type}\")\n            return f\"✓ Converted lists from {from_type} to {to_type}\"
        except Exception as e:
            logger.error(f\"Failed to convert list style: {e}\")
            return f\"✗ Failed to convert list style: {e}\"

    def _set_list_style(self, paragraph, list_type: str):
        """Apply list style to a paragraph."""
        try:
            # Remove existing numbering
            pPr = paragraph._element.pPr
            if pPr is not None:
                numPr = pPr.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                if numPr is not None:
                    pPr.remove(numPr)
            
            # Apply new numbering (simplified - use built-in styles)
            if list_type == 'bulleted':
                paragraph.style = 'List Bullet'
            elif list_type == 'numbered':
                paragraph.style = 'List Number'
        except Exception as e:
            logger.error(f\"Failed to set list style: {e}\")

    def _apply_numbering(self, paragraph, num_id: int, ilvl: int):
        """Apply numbering XML to a paragraph."""
        try:
            pPr = paragraph._element.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            
            ilvl_elem = OxmlElement('w:ilvl')
            ilvl_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(ilvl))
            numPr.append(ilvl_elem)
            
            numId_elem = OxmlElement('w:numId')
            numId_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(num_id))
            numPr.append(numId_elem)
            
            pPr.append(numPr)
        except Exception as e:
            logger.error(f\"Failed to apply numbering: {e}\")

    def remove_paragraphs_by_pattern(self, pattern: str = None, indices: List[int] = None):
        """Remove paragraphs matching a pattern or at specific indices."""
        try:
            removed_count = 0
            paragraphs_to_remove = []
            
            for i, para in enumerate(self.doc.paragraphs):
                should_remove = False
                
                if indices and i in indices:
                    should_remove = True
                elif pattern and pattern.lower() in para.text.lower():
                    should_remove = True
                
                if should_remove:
                    paragraphs_to_remove.append(para)
            
            for para in paragraphs_to_remove:
                p = para._element
                p.getparent().remove(p)
                removed_count += 1
            
            logger.info(f\"Removed {removed_count} paragraphs\")
            return f\"✓ Removed {removed_count} paragraphs\"
        except Exception as e:
            logger.error(f\"Failed to remove paragraphs: {e}\")
            return f\"✗ Failed to remove paragraphs: {e}\"

    def modify_text_formatting(
        self,
        target_text: Optional[str] = None,
        target_style: Optional[str] = None,
        para_indices: Optional[List[int]] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None,
        color: Optional[str] = None,
        font_size: Optional[int] = None
    ):
        \"\"\"Modify text formatting with flexible targeting.\"\"\"
        try:
            modified_count = 0
            
            for i, para in enumerate(self.doc.paragraphs):
                should_modify = False
                
                # Check targeting criteria
                if para_indices and i in para_indices:
                    should_modify = True
                elif target_style and para.style.name == target_style:
                    should_modify = True
                elif target_text and target_text.lower() in para.text.lower():
                    should_modify = True
                elif not para_indices and not target_style and not target_text:
                    should_modify = True  # Apply to all if no targeting specified
                
                if not should_modify:
                    continue
                
                # Apply formatting to all runs
                for run in para.runs:
                    if bold is not None:
                        run.bold = bold
                    if italic is not None:
                        run.italic = italic
                    if underline is not None:
                        run.underline = underline
                    if color:
                        color_rgb = self._parse_color(color)
                        run.font.color.rgb = RGBColor(*color_rgb)
                    if font_size:
                        run.font.size = Pt(font_size)
                
                modified_count += 1
            
            logger.info(f\"Modified formatting for {modified_count} paragraphs\")
            return f\"✓ Modified formatting for {modified_count} paragraphs\"
        except Exception as e:
            logger.error(f\"Failed to modify formatting: {e}\")
            return f\"✗ Failed to modify formatting: {e}\"

    def remove_all_formatting(self, para_indices: Optional[List[int]] = None):
        \"\"\"Remove all formatting from specified paragraphs or entire document.\"\"\"
        try:
            for i, para in enumerate(self.doc.paragraphs):
                if para_indices and i not in para_indices:
                    continue
                
                for run in para.runs:
                    run.bold = False
                    run.italic = False
                    run.underline = False
                    run.font.color.rgb = None
                    run.font.size = None
            
            logger.info(\"Removed all formatting\")
            return \"✓ Removed all formatting\"
        except Exception as e:
            logger.error(f\"Failed to remove formatting: {e}\")
            return f\"✗ Failed to remove formatting: {e}\"

    def delete_tables(self, table_indices: Optional[List[int]] = None):
        \"\"\"Delete specific tables or all tables.\"\"\"
        try:
            if table_indices is None:
                table_indices = list(range(len(self.doc.tables)))
            
            # Delete in reverse order to maintain indices
            for i in sorted(table_indices, reverse=True):
                if 0 <= i < len(self.doc.tables):
                    table = self.doc.tables[i]
                    tbl = table._element
                    tbl.getparent().remove(tbl)
            
            logger.info(f\"Deleted {len(table_indices)} tables\")
            return f\"✓ Deleted {len(table_indices)} tables\"
        except Exception as e:
            logger.error(f\"Failed to delete tables: {e}\")
            return f\"✗ Failed to delete tables: {e}\"

    def modify_table(self, table_index: int, modifications: Dict[str, Any]):
        \"\"\"Modify table cells, structure, or formatting.\"\"\"
        try:
            if table_index >= len(self.doc.tables):
                return \"✗ Invalid table index\"
            
            table = self.doc.tables[table_index]
            
            # Apply modifications
            if 'cell_value' in modifications:
                row = modifications.get('row', 0)
                col = modifications.get('col', 0)
                value = modifications['cell_value']
                if row < len(table.rows) and col < len(table.columns):
                    table.cell(row, col).text = str(value)
            
            if 'add_row' in modifications:
                cells = table.add_row().cells
                if 'row_values' in modifications:
                    for i, val in enumerate(modifications['row_values']):
                        if i < len(cells):
                            cells[i].text = str(val)
            
            if 'add_column' in modifications:
                table.add_column(modifications.get('column_width', Inches(1.5)))
            
            logger.info(f\"Modified table {table_index}\")
            return f\"✓ Modified table {table_index}\"
        except Exception as e:
            logger.error(f\"Failed to modify table: {e}\")
            return f\"✗ Failed to modify table: {e}\"

    def change_style(self, para_indices: List[int], new_style: str):
        \"\"\"Change the style of specified paragraphs.\"\"\"
        try:
            mapped_style = self._map_style(new_style)
            for i in para_indices:
                if 0 <= i < len(self.doc.paragraphs):
                    self.doc.paragraphs[i].style = mapped_style
            
            logger.info(f\"Changed style for {len(para_indices)} paragraphs\")
            return f\"✓ Changed style to {mapped_style}\"
        except Exception as e:
            logger.error(f\"Failed to change style: {e}\")
            return f\"✗ Failed to change style: {e}\"

    def remove_all_lists(self):
        \"\"\"Remove all list formatting from the document.\"\"\"
        try:
            for para in self.doc.paragraphs:
                if self._is_list_item(para):
                    para.style = 'Normal'
                    # Remove numbering
                    pPr = para._element.pPr
                    if pPr is not None:
                        numPr = pPr.find('.//w:numPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if numPr is not None:
                            pPr.remove(numPr)
            
            logger.info(\"Removed all lists\")
            return \"✓ Removed all lists\"
        except Exception as e:
            logger.error(f\"Failed to remove lists: {e}\")
            return f\"✗ Failed to remove lists: {e}\"

    def clear_document_content(self, keep_structure: bool = False):
        \"\"\"Clear document content, optionally keeping structure.\"\"\"
        try:
            if keep_structure:
                # Keep headings and structure, clear only body text
                for para in self.doc.paragraphs:
                    if 'Heading' not in para.style.name and para.style.name != 'Title':
                        para.text = ''
            else:
                # Clear everything
                for element in self.doc.element.body:
                    self.doc.element.body.remove(element)
            
            logger.info(\"Cleared document content\")
            return \"✓ Document content cleared\"
        except Exception as e:
            logger.error(f\"Failed to clear content: {e}\")
            return f\"✗ Failed to clear content: {e}\"

    def find_and_delete(self, search_text: str, case_sensitive: bool = False):
        \"\"\"Find and delete paragraphs containing specific text.\"\"\"
        try:
            deleted_count = 0
            paragraphs_to_delete = []
            
            for para in self.doc.paragraphs:
                text = para.text if case_sensitive else para.text.lower()
                search = search_text if case_sensitive else search_text.lower()
                
                if search in text:
                    paragraphs_to_delete.append(para)
            
            for para in paragraphs_to_delete:
                p = para._element
                p.getparent().remove(p)
                deleted_count += 1
            
            logger.info(f\"Deleted {deleted_count} paragraphs\")
            return f\"✓ Deleted {deleted_count} paragraphs containing '{search_text}'\"
        except Exception as e:
            logger.error(f\"Failed to find and delete: {e}\")
            return f\"✗ Failed to find and delete: {e}\"

    def replace_all_text(self, old_text: str, new_text: str, case_sensitive: bool = False):
        \"\"\"Replace all occurrences of text in the document.\"\"\"
        try:
            replaced_count = 0
            
            for para in self.doc.paragraphs:
                if case_sensitive:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)
                        replaced_count += 1
                else:
                    if old_text.lower() in para.text.lower():
                        # Case-insensitive replacement
                        import re
                        para.text = re.sub(re.escape(old_text), new_text, para.text, flags=re.IGNORECASE)
                        replaced_count += 1
            
            # Also check tables
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if case_sensitive:
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)
                                replaced_count += 1
                        else:
                            if old_text.lower() in cell.text.lower():
                                import re
                                cell.text = re.sub(re.escape(old_text), new_text, cell.text, flags=re.IGNORECASE)
                                replaced_count += 1
            
            logger.info(f\"Replaced text in {replaced_count} locations\")
            return f\"✓ Replaced '{old_text}' with '{new_text}' in {replaced_count} locations\"
        except Exception as e:
            logger.error(f\"Failed to replace text: {e}\")
            return f\"✗ Failed to replace text: {e}\"

    def remove_empty_paragraphs(self):
        \"\"\"Remove all empty paragraphs from the document.\"\"\"
        try:
            removed_count = 0
            paragraphs_to_remove = []
            
            for para in self.doc.paragraphs:
                if not para.text.strip():
                    paragraphs_to_remove.append(para)
            
            for para in paragraphs_to_remove:
                p = para._element
                p.getparent().remove(p)
                removed_count += 1
            
            logger.info(f\"Removed {removed_count} empty paragraphs\")
            return f\"✓ Removed {removed_count} empty paragraphs\"
        except Exception as e:
            logger.error(f\"Failed to remove empty paragraphs: {e}\")
            return f\"✗ Failed to remove empty paragraphs: {e}\"

    def detect_pseudo_lists(self) -> List[Dict[str, Any]]:
        \"\"\"Detect text that looks like lists but isn't formatted as such.\"\"\"
        try:
            pseudo_lists = []
            current_list = None
            
            for i, para in enumerate(self.doc.paragraphs):
                if self._is_list_item(para):
                    current_list = None
                    continue
                
                text = para.text.strip()
                if not text:
                    current_list = None
                    continue
                
                # Check if starts with list indicators
                is_numbered = text and text[0].isdigit() and len(text) > 1 and text[1] in ['.', ')', ':']
                is_bulleted = text and text[0] in ['•', '-', '*', '·', '○']
                
                if is_numbered or is_bulleted:
                    if current_list is None:
                        current_list = {
                            'type': 'numbered' if is_numbered else 'bulleted',
                            'start_index': i,
                            'paragraphs': []
                        }
                        pseudo_lists.append(current_list)
                    
                    current_list['paragraphs'].append({
                        'index': i,
                        'text': text[:100]
                    })
                    current_list['end_index'] = i
                else:
                    current_list = None
            
            logger.info(f\"Detected {len(pseudo_lists)} pseudo-lists\")
            return pseudo_lists
        except Exception as e:
            logger.error(f\"Failed to detect pseudo-lists: {e}\")
            return []

    def convert_text_to_list(self, para_indices: Optional[List[int]] = None, list_type: str = 'numbered'):
        \"\"\"Convert plain text paragraphs to proper lists.\"\"\"
        try:
            if para_indices is None:
                # Detect pseudo-lists and convert them
                pseudo_lists = self.detect_pseudo_lists()
                if not pseudo_lists:
                    return \"✗ No pseudo-lists detected\"
                
                converted_count = 0
                for pseudo_list in pseudo_lists:
                    for item in pseudo_list['paragraphs']:
                        idx = item['index']
                        para = self.doc.paragraphs[idx]
                        
                        # Remove list indicator from text
                        text = para.text.strip()
                        if text[0].isdigit():
                            # Remove number and separator
                            text = text.split('.', 1)[1].strip() if '.' in text else text[1:].strip()
                        elif text[0] in ['•', '-', '*', '·', '○']:
                            text = text[1:].strip()
                        
                        para.text = text
                        self._set_list_style(para, pseudo_list['type'])
                        converted_count += 1
                
                logger.info(f\"Converted {converted_count} pseudo-list items\")
                return f\"✓ Converted {converted_count} items to lists\"
            else:
                # Convert specific paragraphs to lists
                for idx in para_indices:
                    if 0 <= idx < len(self.doc.paragraphs):
                        para = self.doc.paragraphs[idx]
                        self._set_list_style(para, list_type)
                
                logger.info(f\"Converted {len(para_indices)} paragraphs to {list_type} list\")
                return f\"✓ Converted {len(para_indices)} paragraphs to {list_type} list\"
        except Exception as e:
            logger.error(f\"Failed to convert to list: {e}\")
            return f\"✗ Failed to convert to list: {e}\"

    def move_paragraph(self, from_index: int, to_index: int):
        \"\"\"Move a paragraph from one position to another.\"\"\"
        try:
            if not (0 <= from_index < len(self.doc.paragraphs)):
                return \"✗ Invalid from_index\"
            if not (0 <= to_index <= len(self.doc.paragraphs)):
                return \"✗ Invalid to_index\"
            
            # Get paragraph element
            para_element = self.doc.paragraphs[from_index]._element
            parent = para_element.getparent()
            
            # Remove from current position
            parent.remove(para_element)
            
            # Insert at new position
            if to_index >= len(self.doc.paragraphs):
                parent.append(para_element)
            else:
                ref_element = self.doc.paragraphs[to_index]._element
                parent.insert(parent.index(ref_element), para_element)
            
            logger.info(f\"Moved paragraph from {from_index} to {to_index}\")
            return f\"✓ Moved paragraph from position {from_index} to {to_index}\"
        except Exception as e:
            logger.error(f\"Failed to move paragraph: {e}\")
            return f\"✗ Failed to move paragraph: {e}\"

    def insert_paragraph_at(self, index: int, text: str, style: str = 'Normal', **formatting):
        \"\"\"Insert a new paragraph at a specific position with formatting.\"\"\"
        try:
            if not (0 <= index <= len(self.doc.paragraphs)):
                return \"✗ Invalid index\"
            
            mapped_style = self._map_style(style)
            
            # Create new paragraph element
            new_para = self.doc.add_paragraph(text, style=mapped_style)
            
            # Apply formatting if provided
            if formatting:
                for run in new_para.runs:
                    if 'bold' in formatting:
                        run.bold = formatting['bold']
                    if 'italic' in formatting:
                        run.italic = formatting['italic']
                    if 'underline' in formatting:
                        run.underline = formatting['underline']
                    if 'color' in formatting:
                        color_rgb = self._parse_color(formatting['color'])
                        run.font.color.rgb = RGBColor(*color_rgb)
                    if 'font_size' in formatting:
                        run.font.size = Pt(formatting['font_size'])
            
            # Move to correct position if not at end
            if index < len(self.doc.paragraphs) - 1:
                para_element = new_para._element
                parent = para_element.getparent()
                parent.remove(para_element)
                
                ref_element = self.doc.paragraphs[index]._element
                parent.insert(parent.index(ref_element), para_element)
            
            logger.info(f\"Inserted paragraph at position {index}\")
            return f\"✓ Inserted paragraph at position {index}\"
        except Exception as e:
            logger.error(f\"Failed to insert paragraph: {e}\")
            return f\"✗ Failed to insert paragraph: {e}\"
