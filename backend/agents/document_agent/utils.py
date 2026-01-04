"""
Document Agent - Utilities

Helper functions for document processing, conversion, and display.
Optimized for cloud deployment with minimal dependencies.
"""

import os
import base64
import logging
from pathlib import Path
from typing import Dict, Any, Tuple
import mimetypes

logger = logging.getLogger(__name__)


def extract_document_content(file_path: str) -> Tuple[str, str]:
    """
    Extract text content from various document formats.
    Returns (content, file_type)
    """
    file_ext = Path(file_path).suffix.lower()

    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read(), 'text'

        elif file_ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
            return content, 'docx'

        elif file_ext == '.pdf':
            from pypdf import PdfReader
            content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            return '\n'.join(content), 'pdf'

        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    except Exception as e:
        logger.error(f"Failed to extract content from {file_path}: {e}")
        raise


def create_docx(content: str, file_path: str) -> None:
    """Create a Word document from text content."""
    try:
        from docx import Document
        doc = Document()
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        doc.save(file_path)
        logger.info(f"Created DOCX: {file_path}")
    except Exception as e:
        logger.error(f"Failed to create DOCX: {e}")
        raise


def create_pdf(content: str, file_path: str) -> None:
    """Create a PDF document from text content."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for paragraph in content.split('\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, styles['Normal']))
                story.append(Spacer(1, 12))

        doc.build(story)
        logger.info(f"Created PDF: {file_path}")
    except Exception as e:
        logger.error(f"Failed to create PDF: {e}")
        raise


def get_file_base64(file_path: str) -> str:
    """Convert file to base64 for display."""
    try:
        with open(file_path, 'rb') as f:
            return base64.b64encode(f.read()).decode('utf-8')
    except Exception as e:
        logger.error(f"Failed to encode file: {e}")
        raise


def create_canvas_display(
    canvas_type: str,
    canvas_data: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Create canvas display data for frontend rendering.
    Supports: pdf, docx, text, image
    """
    base_display = {
        'canvas_type': canvas_type,
        'timestamp': int(__import__('time').time() * 1000),
        'no_cache': True
    }
    base_display.update(canvas_data)
    return base_display


def create_pdf_canvas_display(
    file_path: str,
    title: str,
    original_type: str = 'pdf'
) -> Dict[str, Any]:
    """Create PDF canvas display with cache busting."""
    pdf_base64 = get_file_base64(file_path)
    return create_canvas_display(
        canvas_type='pdf',
        canvas_data={
            'title': title,
            'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
            'file_path': file_path,
            'original_type': original_type
        }
    )


def analyze_document_structure(file_path: str) -> Dict[str, Any]:
    """
    Analyze document structure for context-aware editing.
    Returns styles, headings, tables, and statistics.
    """
    try:
        from docx import Document
        doc = Document(file_path)

        styles_used = {}
        for para in doc.paragraphs:
            style_name = para.style.name
            styles_used[style_name] = styles_used.get(style_name, 0) + 1

        headings = []
        for para in doc.paragraphs:
            if 'Heading' in para.style.name:
                headings.append({
                    'level': para.style.name,
                    'text': para.text[:100]
                })

        return {
            'styles_used': styles_used,
            'headings': headings[:10],
            'table_count': len(doc.tables),
            'total_paragraphs': len(doc.paragraphs),
            'total_sections': len(doc.sections),
            'file_name': Path(file_path).name,
            'file_size_kb': Path(file_path).stat().st_size / 1024
        }
    except Exception as e:
        logger.error(f"Failed to analyze document structure: {e}")
        return {
            'error': str(e),
            'file_name': Path(file_path).name
        }


def convert_docx_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    """
    Convert DOCX to PDF for display.
    Uses platform-specific tools (LibreOffice on Linux/Mac, docx2pdf on Windows).
    """
    try:
        import platform
        import time
        from pathlib import Path

        if pdf_path is None:
            timestamp = int(time.time() * 1000)
            base_name = Path(docx_path).stem
            pdf_path = str(Path(docx_path).parent / f"{base_name}_display_{timestamp}.pdf")

        if platform.system() == 'Windows':
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
            except ImportError:
                logger.warning("docx2pdf not available on Windows")
                return docx_path  # Return original

        else:
            # Linux/Mac: use LibreOffice
            import subprocess
            subprocess.run(
                [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', str(Path(pdf_path).parent), docx_path
                ],
                check=True,
                capture_output=True
            )

            generated_pdf = str(Path(docx_path).with_suffix('.pdf'))
            if generated_pdf != pdf_path and os.path.exists(generated_pdf):
                os.rename(generated_pdf, pdf_path)

        logger.info(f"Converted DOCX to PDF: {pdf_path}")
        return pdf_path

    except Exception as e:
        logger.error(f"Failed to convert DOCX to PDF: {e}")
        return docx_path  # Fallback to original


def get_mime_type(file_path: str) -> str:
    """Get MIME type for file."""
    mime_type, _ = mimetypes.guess_type(file_path)
    return mime_type or 'application/octet-stream'


def ensure_directory(dir_path: str) -> Path:
    """Ensure directory exists, create if necessary."""
    path = Path(dir_path)
    path.mkdir(parents=True, exist_ok=True)
    return path
