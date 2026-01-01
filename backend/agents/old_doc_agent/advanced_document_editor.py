# agents/advanced_document_editor.py
# Advanced document editing capabilities for Word documents

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm, Twips, Emu
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH, 
    WD_LINE_SPACING, 
    WD_BREAK,
    WD_UNDERLINE,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER
)
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
from typing import Dict, Any, List, Optional, Tuple
import os

logger = logging.getLogger(__name__)

class AdvancedDocumentEditor:
    """Advanced document editing with comprehensive Word capabilities."""
    
    # Color mapping
    COLORS = {
        'red': (255, 0, 0),
        'blue': (0, 0, 255),
        'green': (0, 128, 0),
        'yellow': (255, 255, 0),
        'orange': (255, 165, 0),
        'purple': (128, 0, 128),
        'black': (0, 0, 0),
        'white': (255, 255, 255),
        'gray': (128, 128, 128),
        'grey': (128, 128, 128),
        'cyan': (0, 255, 255),
        'magenta': (255, 0, 255),
        'brown': (165, 42, 42),
        'pink': (255, 192, 203),
        'lime': (0, 255, 0),
        'navy': (0, 0, 128),
        'teal': (0, 128, 128),
        'maroon': (128, 0, 0),
        'olive': (128, 128, 0)
    }
    
    # Style mapping
    STYLE_MAPPING = {
        'subheading': 'Heading 2',
        'subheadings': 'Heading 2',
        'sub-heading': 'Heading 2',
        'subtitle': 'Heading 2',
        'heading': 'Heading 1',
        'headings': 'Heading 1',
        'title': 'Title',
        'heading1': 'Heading 1',
        'heading2': 'Heading 2',
        'heading3': 'Heading 3',
        'h1': 'Heading 1',
        'h2': 'Heading 2',
        'h3': 'Heading 3',
        'normal': 'Normal',
        'body': 'Normal'
    }
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
    
    def save(self):
        """Save the document."""
        self.doc.save(self.file_path)
        logger.info(f"Document saved: {self.file_path}")
    
    def _map_style(self, style: str) -> str:
        """Map common style names to Word style names."""
        return self.STYLE_MAPPING.get(style.lower(), style)
    
    def _parse_color(self, color: Any) -> Tuple[int, int, int]:
        """Parse color from string or RGB tuple."""
        if isinstance(color, str):
            return self.COLORS.get(color.lower(), (0, 0, 0))
        elif isinstance(color, (list, tuple)) and len(color) == 3:
            return tuple(color)
        return (0, 0, 0)
    
    # ==================== TEXT FORMATTING ====================
    
    def format_text(self, params: Dict[str, Any]) -> str:
        """
        Format text with various options.
        Supports: style, text, bold, italic, underline, font_size, color, highlight
        """
        text_to_format = params.get('text', '')
        style_to_format = params.get('style', '')
        bold = params.get('bold', False)
        italic = params.get('italic', False)
        underline = params.get('underline', False)
        strikethrough = params.get('strikethrough', False)
        font_size = params.get('font_size', None)
        color = params.get('color', None)
        highlight = params.get('highlight', None)
        font_name = params.get('font', None)
        
        formatted_count = 0
        
        # Format by style
        if style_to_format:
            mapped_style = self._map_style(style_to_format)
            logger.info(f"Formatting paragraphs with style: {mapped_style}")
            
            for para in self.doc.paragraphs:
                if para.style.name == mapped_style or mapped_style.lower() in para.style.name.lower():
                    for run in para.runs:
                        self._apply_run_formatting(run, bold, italic, underline, strikethrough,
                                                   font_size, color, highlight, font_name)
                        formatted_count += 1
        
        # Format by text content
        elif text_to_format:
            for para in self.doc.paragraphs:
                for run in para.runs:
                    if text_to_format in run.text:
                        self._apply_run_formatting(run, bold, italic, underline, strikethrough,
                                                   font_size, color, highlight, font_name)
                        formatted_count += 1
        
        return f"Formatted {formatted_count} text runs"
    
    def _apply_run_formatting(self, run, bold, italic, underline, strikethrough,
                             font_size, color, highlight, font_name):
        """Apply formatting to a run."""
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if strikethrough is not None:
            run.font.strike = strikethrough
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            r, g, b = self._parse_color(color)
            run.font.color.rgb = RGBColor(r, g, b)
        if highlight:
            # Highlight colors: yellow, green, cyan, magenta, blue, red, etc.
            run.font.highlight_color = self._get_highlight_color(highlight)
        if font_name:
            run.font.name = font_name
    
    def _get_highlight_color(self, color_name: str):
        """Get Word highlight color constant."""
        from docx.enum.text import WD_COLOR_INDEX
        color_map = {
            'yellow': WD_COLOR_INDEX.YELLOW,
            'green': WD_COLOR_INDEX.BRIGHT_GREEN,
            'cyan': WD_COLOR_INDEX.TURQUOISE,
            'magenta': WD_COLOR_INDEX.PINK,
            'blue': WD_COLOR_INDEX.BLUE,
            'red': WD_COLOR_INDEX.RED,
            'gray': WD_COLOR_INDEX.GRAY_25,
            'grey': WD_COLOR_INDEX.GRAY_25
        }
        return color_map.get(color_name.lower(), WD_COLOR_INDEX.YELLOW)
    
    # ==================== PARAGRAPH OPERATIONS ====================
    
    def add_paragraph(self, params: Dict[str, Any]) -> str:
        """Add a new paragraph with formatting."""
        text = params.get('text', '')
        style = params.get('style', None)
        alignment = params.get('alignment', 'left')
        bold = params.get('bold', False)
        italic = params.get('italic', False)
        font_size = params.get('font_size', None)
        color = params.get('color', None)
        spacing_before = params.get('spacing_before', None)
        spacing_after = params.get('spacing_after', None)
        line_spacing = params.get('line_spacing', None)
        
        para = self.doc.add_paragraph(text, style=style)
        
        # Set alignment
        if alignment:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            para.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
        
        # Apply formatting to runs
        for run in para.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if font_size:
                run.font.size = Pt(font_size)
            if color:
                r, g, b = self._parse_color(color)
                run.font.color.rgb = RGBColor(r, g, b)
        
        # Set spacing
        if spacing_before:
            para.paragraph_format.space_before = Pt(spacing_before)
        if spacing_after:
            para.paragraph_format.space_after = Pt(spacing_after)
        if line_spacing:
            para.paragraph_format.line_spacing = line_spacing
        
        return f"Added paragraph: {text[:50]}..."
    
    def remove_paragraph(self, params: Dict[str, Any]) -> str:
        """Remove paragraph by index or text content."""
        index = params.get('index', None)
        text = params.get('text', None)
        
        if index is not None:
            if 0 <= index < len(self.doc.paragraphs):
                p = self.doc.paragraphs[index]
                p._element.getparent().remove(p._element)
                return f"Removed paragraph at index {index}"
        
        if text:
            removed_count = 0
            for para in list(self.doc.paragraphs):
                if text in para.text:
                    para._element.getparent().remove(para._element)
                    removed_count += 1
            return f"Removed {removed_count} paragraphs containing '{text}'"
        
        return "No paragraphs removed"
    
    def replace_text(self, params: Dict[str, Any]) -> str:
        """Replace text throughout the document."""
        old_text = params.get('old_text', '')
        new_text = params.get('new_text', '')
        case_sensitive = params.get('case_sensitive', True)
        
        replaced_count = 0
        
        for para in self.doc.paragraphs:
            for run in para.runs:
                if case_sensitive:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replaced_count += 1
                else:
                    if old_text.lower() in run.text.lower():
                        # Case-insensitive replacement
                        import re
                        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                        run.text = pattern.sub(new_text, run.text)
                        replaced_count += 1
        
        return f"Replaced {replaced_count} occurrences of '{old_text}' with '{new_text}'"
    
    # ==================== LISTS ====================
    
    def add_bullet_list(self, params: Dict[str, Any]) -> str:
        """Add a bulleted list."""
        items = params.get('items', [])
        style = params.get('style', 'List Bullet')
        
        for item in items:
            self.doc.add_paragraph(item, style=style)
        
        return f"Added bullet list with {len(items)} items"
    
    def add_numbered_list(self, params: Dict[str, Any]) -> str:
        """Add a numbered list."""
        items = params.get('items', [])
        style = params.get('style', 'List Number')
        
        for item in items:
            self.doc.add_paragraph(item, style=style)
        
        return f"Added numbered list with {len(items)} items"
    
    # ==================== HEADINGS ====================
    
    def add_heading(self, params: Dict[str, Any]) -> str:
        """Add a heading."""
        text = params.get('text', '')
        level = params.get('level', 1)
        
        self.doc.add_heading(text, level=level)
        return f"Added heading level {level}: {text}"
    
    # ==================== TABLES ====================
    
    def add_table(self, params: Dict[str, Any]) -> str:
        """Add a table with data."""
        rows = params.get('rows', 2)
        cols = params.get('cols', 2)
        data = params.get('data', [])
        style = params.get('style', 'Table Grid')
        header_row = params.get('header_row', True)
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style
        
        # Populate table with data
        if data:
            for i, row_data in enumerate(data):
                if i < len(table.rows):
                    for j, cell_value in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = str(cell_value)
        
        # Make first row bold if header
        if header_row and len(table.rows) > 0:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        return f"Added table with {rows} rows and {cols} columns"
    
    def modify_table(self, params: Dict[str, Any]) -> str:
        """Modify an existing table."""
        table_index = params.get('table_index', 0)
        action = params.get('action', 'add_row')  # add_row, add_column, delete_row, delete_column
        
        if table_index >= len(self.doc.tables):
            return f"Table index {table_index} out of range"
        
        table = self.doc.tables[table_index]
        
        if action == 'add_row':
            table.add_row()
            return f"Added row to table {table_index}"
        elif action == 'add_column':
            table.add_column(Inches(1.0))
            return f"Added column to table {table_index}"
        # Note: Deleting rows/columns is complex in python-docx
        
        return f"Modified table {table_index}"
    
    # ==================== IMAGES ====================
    
    def add_image(self, params: Dict[str, Any]) -> str:
        """Add an image to the document."""
        image_path = params.get('image_path')
        width = params.get('width', 4)  # Width in inches
        height = params.get('height', None)  # Height in inches (optional)
        caption = params.get('caption', None)
        
        if not image_path or not os.path.exists(image_path):
            return f"Image not found: {image_path}"
        
        # Add image
        if height:
            self.doc.add_picture(image_path, width=Inches(width), height=Inches(height))
        else:
            self.doc.add_picture(image_path, width=Inches(width))
        
        # Add caption if provided
        if caption:
            self.doc.add_paragraph(caption, style='Caption')
        
        return f"Added image: {os.path.basename(image_path)}"
    
    # ==================== PAGE BREAKS ====================
    
    def add_page_break(self, params: Dict[str, Any]) -> str:
        """Add a page break."""
        self.doc.add_page_break()
        return "Added page break"
    
    # ==================== SECTIONS ====================
    
    def add_section(self, params: Dict[str, Any]) -> str:
        """Add a new section."""
        from docx.enum.section import WD_SECTION
        
        section_type = params.get('type', 'new_page')  # new_page, continuous, even_page, odd_page
        
        section_map = {
            'new_page': WD_SECTION.NEW_PAGE,
            'continuous': WD_SECTION.CONTINUOUS,
            'even_page': WD_SECTION.EVEN_PAGE,
            'odd_page': WD_SECTION.ODD_PAGE
        }
        
        section = self.doc.add_section(section_map.get(section_type, WD_SECTION.NEW_PAGE))
        return f"Added section: {section_type}"
    
    # ==================== HEADERS & FOOTERS ====================
    
    def add_header(self, params: Dict[str, Any]) -> str:
        """Add text to header."""
        text = params.get('text', '')
        alignment = params.get('alignment', 'center')
        
        section = self.doc.sections[0]
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = text
        
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        para.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
        
        return f"Added header: {text}"
    
    def add_footer(self, params: Dict[str, Any]) -> str:
        """Add text to footer."""
        text = params.get('text', '')
        alignment = params.get('alignment', 'center')
        page_numbers = params.get('page_numbers', False)
        
        section = self.doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        
        if page_numbers:
            para.text = f"{text} - Page "
            # Add page number field
            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)
        else:
            para.text = text
        
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        para.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
        
        return f"Added footer: {text}"
    
    # ==================== HYPERLINKS ====================
    
    def add_hyperlink(self, params: Dict[str, Any]) -> str:
        """Add a hyperlink to text."""
        text = params.get('text', '')
        url = params.get('url', '')
        paragraph_index = params.get('paragraph_index', -1)  # -1 means add new paragraph
        
        if paragraph_index == -1:
            para = self.doc.add_paragraph()
        else:
            if paragraph_index < len(self.doc.paragraphs):
                para = self.doc.paragraphs[paragraph_index]
            else:
                return f"Paragraph index {paragraph_index} out of range"
        
        # Add hyperlink
        self._add_hyperlink(para, url, text)
        
        return f"Added hyperlink: {text} -> {url}"
    
    def _add_hyperlink(self, paragraph, url, text):
        """Helper to add hyperlink to paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Style the hyperlink (blue and underlined)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
    
    # ==================== COMMENTS ====================
    
    def add_comment(self, params: Dict[str, Any]) -> str:
        """Add a comment (note: limited support in python-docx)."""
        text = params.get('text', '')
        # Comments require complex XML manipulation
        # For now, add as a footnote-style paragraph
        para = self.doc.add_paragraph()
        run = para.add_run(f"[Comment: {text}]")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        return f"Added comment: {text}"
    
    # ==================== FIND & REPLACE ====================
    
    def find_and_replace_with_format(self, params: Dict[str, Any]) -> str:
        """Find text and replace with formatted text."""
        find_text = params.get('find_text', '')
        replace_text = params.get('replace_text', '')
        bold = params.get('bold', False)
        italic = params.get('italic', False)
        color = params.get('color', None)
        
        replaced_count = 0
        
        for para in self.doc.paragraphs:
            if find_text in para.text:
                # Replace and format
                for run in para.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        if bold:
                            run.bold = True
                        if italic:
                            run.italic = True
                        if color:
                            r, g, b = self._parse_color(color)
                            run.font.color.rgb = RGBColor(r, g, b)
                        replaced_count += 1
        
        return f"Found and replaced {replaced_count} occurrences"
    
    # ==================== ADVANCED PARAGRAPH FORMATTING ====================
    
    def set_paragraph_spacing(self, params: Dict[str, Any]) -> str:
        """Set advanced paragraph spacing."""
        paragraph_index = params.get('paragraph_index', -1)
        line_spacing = params.get('line_spacing', None)
        line_spacing_rule = params.get('line_spacing_rule', 'single')  # single, double, multiple, exactly, at_least
        space_before = params.get('space_before', None)
        space_after = params.get('space_after', None)
        
        if paragraph_index == -1:
            # Apply to all paragraphs
            paragraphs = self.doc.paragraphs
        else:
            if paragraph_index >= len(self.doc.paragraphs):
                return f"Paragraph index {paragraph_index} out of range"
            paragraphs = [self.doc.paragraphs[paragraph_index]]
        
        for para in paragraphs:
            if space_before:
                para.paragraph_format.space_before = Pt(space_before)
            if space_after:
                para.paragraph_format.space_after = Pt(space_after)
            if line_spacing:
                para.paragraph_format.line_spacing = line_spacing
                
                # Set line spacing rule
                rule_map = {
                    'single': WD_LINE_SPACING.SINGLE,
                    'double': WD_LINE_SPACING.DOUBLE,
                    'multiple': WD_LINE_SPACING.MULTIPLE,
                    'exactly': WD_LINE_SPACING.EXACTLY,
                    'at_least': WD_LINE_SPACING.AT_LEAST
                }
                para.paragraph_format.line_spacing_rule = rule_map.get(line_spacing_rule.lower(), WD_LINE_SPACING.SINGLE)
        
        return f"Set spacing for {len(paragraphs)} paragraph(s)"
    
    def set_paragraph_indentation(self, params: Dict[str, Any]) -> str:
        """Set paragraph indentation."""
        paragraph_index = params.get('paragraph_index', -1)
        left_indent = params.get('left_indent', None)
        right_indent = params.get('right_indent', None)
        first_line_indent = params.get('first_line_indent', None)
        
        if paragraph_index == -1:
            paragraphs = self.doc.paragraphs
        else:
            if paragraph_index >= len(self.doc.paragraphs):
                return f"Paragraph index {paragraph_index} out of range"
            paragraphs = [self.doc.paragraphs[paragraph_index]]
        
        for para in paragraphs:
            if left_indent is not None:
                para.paragraph_format.left_indent = Inches(left_indent)
            if right_indent is not None:
                para.paragraph_format.right_indent = Inches(right_indent)
            if first_line_indent is not None:
                para.paragraph_format.first_line_indent = Inches(first_line_indent)
        
        return f"Set indentation for {len(paragraphs)} paragraph(s)"
    
    def set_keep_together(self, params: Dict[str, Any]) -> str:
        """Keep paragraph lines together or with next paragraph."""
        paragraph_index = params.get('paragraph_index', -1)
        keep_together = params.get('keep_together', True)
        keep_with_next = params.get('keep_with_next', False)
        page_break_before = params.get('page_break_before', False)
        widow_control = params.get('widow_control', True)
        
        if paragraph_index == -1:
            paragraphs = self.doc.paragraphs
        else:
            if paragraph_index >= len(self.doc.paragraphs):
                return f"Paragraph index {paragraph_index} out of range"
            paragraphs = [self.doc.paragraphs[paragraph_index]]
        
        for para in paragraphs:
            para.paragraph_format.keep_together = keep_together
            para.paragraph_format.keep_with_next = keep_with_next
            para.paragraph_format.page_break_before = page_break_before
            para.paragraph_format.widow_control = widow_control
        
        return f"Set keep options for {len(paragraphs)} paragraph(s)"
    
    # ==================== TAB STOPS ====================
    
    def add_tab_stops(self, params: Dict[str, Any]) -> str:
        """Add custom tab stops to paragraphs."""
        paragraph_index = params.get('paragraph_index', -1)
        tab_stops = params.get('tab_stops', [])  # List of {position, alignment, leader}
        clear_existing = params.get('clear_existing', False)
        
        if paragraph_index == -1:
            paragraphs = self.doc.paragraphs
        else:
            if paragraph_index >= len(self.doc.paragraphs):
                return f"Paragraph index {paragraph_index} out of range"
            paragraphs = [self.doc.paragraphs[paragraph_index]]
        
        alignment_map = {
            'left': WD_TAB_ALIGNMENT.LEFT,
            'center': WD_TAB_ALIGNMENT.CENTER,
            'right': WD_TAB_ALIGNMENT.RIGHT,
            'decimal': WD_TAB_ALIGNMENT.DECIMAL
        }
        
        leader_map = {
            'spaces': WD_TAB_LEADER.SPACES,
            'dots': WD_TAB_LEADER.DOTS,
            'dashes': WD_TAB_LEADER.DASHES,
            'lines': WD_TAB_LEADER.LINES,
            'heavy': WD_TAB_LEADER.HEAVY,
            'middle_dot': WD_TAB_LEADER.MIDDLE_DOT
        }
        
        for para in paragraphs:
            if clear_existing:
                para.paragraph_format.tab_stops.clear_all()
            
            for tab in tab_stops:
                position = Inches(tab.get('position', 1.0))
                alignment = alignment_map.get(tab.get('alignment', 'left').lower(), WD_TAB_ALIGNMENT.LEFT)
                leader = leader_map.get(tab.get('leader', 'spaces').lower(), WD_TAB_LEADER.SPACES)
                para.paragraph_format.tab_stops.add_tab_stop(position, alignment, leader)
        
        return f"Added {len(tab_stops)} tab stop(s) to {len(paragraphs)} paragraph(s)"
    
    # ==================== ADVANCED FONT FORMATTING ====================
    
    def set_advanced_font(self, params: Dict[str, Any]) -> str:
        """Set advanced font properties."""
        text = params.get('text', '')
        style = params.get('style', '')
        
        # Font properties
        all_caps = params.get('all_caps', None)
        small_caps = params.get('small_caps', None)
        subscript = params.get('subscript', None)
        superscript = params.get('superscript', None)
        double_strike = params.get('double_strike', None)
        emboss = params.get('emboss', None)
        outline = params.get('outline', None)
        shadow = params.get('shadow', None)
        hidden = params.get('hidden', None)
        
        formatted_count = 0
        
        # Find paragraphs to format
        if style:
            mapped_style = self._map_style(style)
            paragraphs = [p for p in self.doc.paragraphs 
                         if p.style.name == mapped_style or mapped_style.lower() in p.style.name.lower()]
        elif text:
            paragraphs = [p for p in self.doc.paragraphs if text in p.text]
        else:
            paragraphs = self.doc.paragraphs
        
        for para in paragraphs:
            for run in para.runs:
                if not text or text in run.text:
                    if all_caps is not None:
                        run.font.all_caps = all_caps
                    if small_caps is not None:
                        run.font.small_caps = small_caps
                    if subscript is not None:
                        run.font.subscript = subscript
                    if superscript is not None:
                        run.font.superscript = superscript
                    if double_strike is not None:
                        run.font.double_strike = double_strike
                    if emboss is not None:
                        run.font.emboss = emboss
                    if outline is not None:
                        run.font.outline = outline
                    if shadow is not None:
                        run.font.shadow = shadow
                    if hidden is not None:
                        run.font.hidden = hidden
                    formatted_count += 1
        
        return f"Applied advanced font formatting to {formatted_count} run(s)"
    
    # ==================== SECTION MANAGEMENT ====================
    
    def set_page_size(self, params: Dict[str, Any]) -> str:
        """Set page size and orientation."""
        section_index = params.get('section_index', 0)
        width = params.get('width', None)  # in inches
        height = params.get('height', None)  # in inches
        orientation = params.get('orientation', None)  # portrait, landscape
        
        if section_index >= len(self.doc.sections):
            return f"Section index {section_index} out of range"
        
        section = self.doc.sections[section_index]
        
        if orientation:
            if orientation.lower() == 'landscape':
                section.orientation = WD_ORIENTATION.LANDSCAPE
                # Swap width and height for landscape
                section.page_width = Inches(11)
                section.page_height = Inches(8.5)
            else:
                section.orientation = WD_ORIENTATION.PORTRAIT
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
        
        if width:
            section.page_width = Inches(width)
        if height:
            section.page_height = Inches(height)
        
        return f"Set page size for section {section_index}"
    
    def set_margins(self, params: Dict[str, Any]) -> str:
        """Set page margins."""
        section_index = params.get('section_index', 0)
        top = params.get('top', None)
        bottom = params.get('bottom', None)
        left = params.get('left', None)
        right = params.get('right', None)
        gutter = params.get('gutter', None)
        
        if section_index >= len(self.doc.sections):
            return f"Section index {section_index} out of range"
        
        section = self.doc.sections[section_index]
        
        if top is not None:
            section.top_margin = Inches(top)
        if bottom is not None:
            section.bottom_margin = Inches(bottom)
        if left is not None:
            section.left_margin = Inches(left)
        if right is not None:
            section.right_margin = Inches(right)
        if gutter is not None:
            section.gutter = Inches(gutter)
        
        return f"Set margins for section {section_index}"
    
    def set_header_footer_distance(self, params: Dict[str, Any]) -> str:
        """Set distance of header/footer from page edge."""
        section_index = params.get('section_index', 0)
        header_distance = params.get('header_distance', None)
        footer_distance = params.get('footer_distance', None)
        
        if section_index >= len(self.doc.sections):
            return f"Section index {section_index} out of range"
        
        section = self.doc.sections[section_index]
        
        if header_distance is not None:
            section.header_distance = Inches(header_distance)
        if footer_distance is not None:
            section.footer_distance = Inches(footer_distance)
        
        return f"Set header/footer distance for section {section_index}"
    
    def set_different_first_page(self, params: Dict[str, Any]) -> str:
        """Enable different first page header/footer."""
        section_index = params.get('section_index', 0)
        different_first = params.get('different_first', True)
        
        if section_index >= len(self.doc.sections):
            return f"Section index {section_index} out of range"
        
        section = self.doc.sections[section_index]
        section.different_first_page_header_footer = different_first
        
        return f"Set different first page for section {section_index}"
    
    # ==================== TABLE ADVANCED OPERATIONS ====================
    
    def set_table_alignment(self, params: Dict[str, Any]) -> str:
        """Set table alignment."""
        table_index = params.get('table_index', 0)
        alignment = params.get('alignment', 'left')  # left, center, right
        
        if table_index >= len(self.doc.tables):
            return f"Table index {table_index} out of range"
        
        table = self.doc.tables[table_index]
        
        alignment_map = {
            'left': WD_TABLE_ALIGNMENT.LEFT,
            'center': WD_TABLE_ALIGNMENT.CENTER,
            'right': WD_TABLE_ALIGNMENT.RIGHT
        }
        
        table.alignment = alignment_map.get(alignment.lower(), WD_TABLE_ALIGNMENT.LEFT)
        
        return f"Set alignment for table {table_index}"
    
    def set_cell_properties(self, params: Dict[str, Any]) -> str:
        """Set cell properties (merge, vertical alignment, etc.)."""
        table_index = params.get('table_index', 0)
        row = params.get('row', 0)
        col = params.get('col', 0)
        vertical_alignment = params.get('vertical_alignment', None)  # top, center, bottom
        width = params.get('width', None)
        
        if table_index >= len(self.doc.tables):
            return f"Table index {table_index} out of range"
        
        table = self.doc.tables[table_index]
        
        if row >= len(table.rows) or col >= len(table.columns):
            return f"Cell ({row}, {col}) out of range"
        
        cell = table.cell(row, col)
        
        if vertical_alignment:
            alignment_map = {
                'top': WD_CELL_VERTICAL_ALIGNMENT.TOP,
                'center': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                'bottom': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            }
            cell.vertical_alignment = alignment_map.get(vertical_alignment.lower(), WD_CELL_VERTICAL_ALIGNMENT.TOP)
        
        if width:
            cell.width = Inches(width)
        
        return f"Set properties for cell ({row}, {col})"
    
    def merge_cells(self, params: Dict[str, Any]) -> str:
        """Merge table cells."""
        table_index = params.get('table_index', 0)
        start_row = params.get('start_row', 0)
        start_col = params.get('start_col', 0)
        end_row = params.get('end_row', 0)
        end_col = params.get('end_col', 0)
        
        if table_index >= len(self.doc.tables):
            return f"Table index {table_index} out of range"
        
        table = self.doc.tables[table_index]
        
        start_cell = table.cell(start_row, start_col)
        end_cell = table.cell(end_row, end_col)
        start_cell.merge(end_cell)
        
        return f"Merged cells from ({start_row},{start_col}) to ({end_row},{end_col})"
    
    def set_row_height(self, params: Dict[str, Any]) -> str:
        """Set row height."""
        table_index = params.get('table_index', 0)
        row_index = params.get('row_index', 0)
        height = params.get('height', 0.5)  # in inches
        height_rule = params.get('height_rule', 'at_least')  # at_least, exactly, auto
        
        if table_index >= len(self.doc.tables):
            return f"Table index {table_index} out of range"
        
        table = self.doc.tables[table_index]
        
        if row_index >= len(table.rows):
            return f"Row index {row_index} out of range"
        
        row = table.rows[row_index]
        row.height = Inches(height)
        
        rule_map = {
            'at_least': WD_ROW_HEIGHT_RULE.AT_LEAST,
            'exactly': WD_ROW_HEIGHT_RULE.EXACTLY,
            'auto': WD_ROW_HEIGHT_RULE.AUTO
        }
        row.height_rule = rule_map.get(height_rule.lower(), WD_ROW_HEIGHT_RULE.AT_LEAST)
        
        return f"Set height for row {row_index}"
    
    def set_column_width(self, params: Dict[str, Any]) -> str:
        """Set column width."""
        table_index = params.get('table_index', 0)
        col_index = params.get('col_index', 0)
        width = params.get('width', 1.0)  # in inches
        
        if table_index >= len(self.doc.tables):
            return f"Table index {table_index} out of range"
        
        table = self.doc.tables[table_index]
        
        if col_index >= len(table.columns):
            return f"Column index {col_index} out of range"
        
        table.columns[col_index].width = Inches(width)
        
        return f"Set width for column {col_index}"
    
    # ==================== STYLES ====================
    
    def create_custom_style(self, params: Dict[str, Any]) -> str:
        """Create a custom style."""
        name = params.get('name', 'CustomStyle')
        style_type = params.get('style_type', 'paragraph')  # paragraph, character, table
        base_style = params.get('base_style', None)
        
        # Font properties
        font_name = params.get('font_name', None)
        font_size = params.get('font_size', None)
        bold = params.get('bold', None)
        italic = params.get('italic', None)
        color = params.get('color', None)
        
        type_map = {
            'paragraph': WD_STYLE_TYPE.PARAGRAPH,
            'character': WD_STYLE_TYPE.CHARACTER,
            'table': WD_STYLE_TYPE.TABLE
        }
        
        style = self.doc.styles.add_style(name, type_map.get(style_type, WD_STYLE_TYPE.PARAGRAPH))
        
        if base_style:
            try:
                style.base_style = self.doc.styles[base_style]
            except KeyError:
                pass
        
        # Apply font properties
        if hasattr(style, 'font'):
            if font_name:
                style.font.name = font_name
            if font_size:
                style.font.size = Pt(font_size)
            if bold is not None:
                style.font.bold = bold
            if italic is not None:
                style.font.italic = italic
            if color:
                r, g, b = self._parse_color(color)
                style.font.color.rgb = RGBColor(r, g, b)
        
        return f"Created custom style: {name}"
    
    # ==================== CORE PROPERTIES (METADATA) ====================
    
    def set_document_properties(self, params: Dict[str, Any]) -> str:
        """Set document metadata properties."""
        core_props = self.doc.core_properties
        
        if 'author' in params:
            core_props.author = params['author']
        if 'title' in params:
            core_props.title = params['title']
        if 'subject' in params:
            core_props.subject = params['subject']
        if 'keywords' in params:
            core_props.keywords = params['keywords']
        if 'category' in params:
            core_props.category = params['category']
        if 'comments' in params:
            core_props.comments = params['comments']
        
        return "Set document properties"
    
    # ==================== MAIN EDIT FUNCTION ====================
    
    def edit(self, instructions: Dict[str, Any]) -> str:
        """
        Main edit function that routes to specific editing methods.
        """
        action = instructions.get('action')
        params = instructions.get('params', {})
        
        action_map = {
            # Text formatting
            'format_text': self.format_text,
            'set_advanced_font': self.set_advanced_font,
            
            # Paragraphs
            'add_paragraph': self.add_paragraph,
            'remove_paragraph': self.remove_paragraph,
            'replace_text': self.replace_text,
            'set_paragraph_spacing': self.set_paragraph_spacing,
            'set_paragraph_indentation': self.set_paragraph_indentation,
            'set_keep_together': self.set_keep_together,
            
            # Lists
            'add_bullet_list': self.add_bullet_list,
            'add_numbered_list': self.add_numbered_list,
            
            # Headings
            'add_heading': self.add_heading,
            
            # Tables
            'add_table': self.add_table,
            'modify_table': self.modify_table,
            'set_table_alignment': self.set_table_alignment,
            'set_cell_properties': self.set_cell_properties,
            'merge_cells': self.merge_cells,
            'set_row_height': self.set_row_height,
            'set_column_width': self.set_column_width,
            
            # Images
            'add_image': self.add_image,
            
            # Page control
            'add_page_break': self.add_page_break,
            'add_section': self.add_section,
            'set_page_size': self.set_page_size,
            'set_margins': self.set_margins,
            
            # Headers/Footers
            'add_header': self.add_header,
            'add_footer': self.add_footer,
            'set_header_footer_distance': self.set_header_footer_distance,
            'set_different_first_page': self.set_different_first_page,
            
            # Hyperlinks
            'add_hyperlink': self.add_hyperlink,
            
            # Comments
            'add_comment': self.add_comment,
            
            # Find & Replace
            'find_and_replace': self.find_and_replace_with_format,
            
            # Tab stops
            'add_tab_stops': self.add_tab_stops,
            
            # Styles
            'create_custom_style': self.create_custom_style,
            
            # Document properties
            'set_document_properties': self.set_document_properties
        }
        
        if action in action_map:
            result = action_map[action](params)
            self.save()
            return result
        else:
            return f"Unknown action: {action}"
