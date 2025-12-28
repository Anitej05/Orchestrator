# agents/document_analysis_agent.py

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel, Field
import uvicorn
import os
from dotenv import load_dotenv
import logging
import base64
from pathlib import Path
from typing import Optional, Dict, Any, List
import sys
import json

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from agents.canvas_utils import create_canvas_display
from agents.document_version_manager import version_manager
from agents.advanced_document_editor import AdvancedDocumentEditor
from agents.document_session_manager import session_manager
from agents.intelligent_document_editor import IntelligentDocumentEditor

# Lazy imports for heavy dependencies - imported only when needed
# This significantly speeds up agent startup time

# Load environment variables from a .env file at the project root
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Configuration & API Key Check ---
CEREBRAS_API_KEY = os.getenv("CEREBRAS_API_KEY")
if not CEREBRAS_API_KEY:
    raise RuntimeError("CEREBRAS_API_KEY is not set in the environment. The agent cannot start.")

# --- Model and Embeddings Loading (Lazy) ---
# Initialize these as None and load them on first request to speed up startup
hf_embeddings = None
llm = None

def get_embeddings():
    """Lazy load the embeddings model on first use."""
    global hf_embeddings
    if hf_embeddings is None:
        logger.info("Loading HuggingFace embeddings model (first request)...")
        from langchain_huggingface import HuggingFaceEmbeddings
        hf_embeddings = HuggingFaceEmbeddings(model_name='all-mpnet-base-v2')
        logger.info("HuggingFace embeddings model loaded successfully")
    return hf_embeddings

def get_llm():
    """Lazy load the LLM on first use."""
    global llm
    if llm is None:
        logger.info("Initializing Cerebras LLM (first request)...")
        from langchain_cerebras import ChatCerebras
        llm = ChatCerebras(model="gpt-oss-120b")
        logger.info("Cerebras LLM initialized successfully")
    return llm

# --- FastAPI Application ---
app = FastAPI(
    title="Document Analysis Agent",
    description="A RAG-based agent that answers questions about documents.",
    version="1.0.0"
)

# --- Pydantic Models for API Data Validation ---
class AnalyzeDocumentRequest(BaseModel):
    """Defines the expected input for the /analyze endpoint."""
    vector_store_path: Optional[str] = Field(None, description="Path to single FAISS vector store (backward compatible)")
    vector_store_paths: Optional[List[str]] = Field(None, description="Paths to multiple FAISS vector stores for multi-document analysis")
    query: str = Field(..., description="The user's question about the document(s).")

class AnalyzeDocumentResponse(BaseModel):
    """Defines the output format for the /analyze endpoint."""
    answer: str
    canvas_display: Optional[Dict[str, Any]] = None

class DisplayDocumentRequest(BaseModel):
    """Request to display a document in canvas."""
    file_path: str = Field(..., description="Path to the document file (PDF, DOCX, TXT)")
    
class DisplayDocumentResponse(BaseModel):
    """Response with document display data."""
    message: str
    canvas_display: Dict[str, Any]
    file_type: str

class CreateDocumentRequest(BaseModel):
    """Request to create a new document."""
    content: Optional[str] = Field(None, description="Content for the document (optional if using reference_file_path)")
    reference_file_path: Optional[str] = Field(None, description="Path to a source file to transform into a document")
    instruction: Optional[str] = Field(None, description="Natural language instruction for how to create the document using LLM")
    file_name: str = Field(..., description="Name for the document file (e.g., 'report.docx')")
    file_type: str = Field(default="docx", description="Document type: 'docx', 'txt', or 'pdf'")
    output_dir: str = Field(default="storage/documents", description="Directory to save the document")

class CreateDocumentResponse(BaseModel):
    """Response after creating a document."""
    message: str
    file_path: str
    canvas_display: Optional[Dict[str, Any]] = None

class EditDocumentRequest(BaseModel):
    """Request to edit an existing document using natural language instruction."""
    file_path: str = Field(..., description="Path to the document to edit")
    instruction: str = Field(..., description="Natural language instruction describing the desired edit")

class EditDocumentResponse(BaseModel):
    """Response after editing a document."""
    message: str
    file_path: str
    canvas_display: Optional[Dict[str, Any]] = None
    can_undo: bool = False
    can_redo: bool = False

class UndoRedoRequest(BaseModel):
    """Request to undo/redo an edit."""
    file_path: str = Field(..., description="Path to the document")

class UndoRedoResponse(BaseModel):
    """Response after undo/redo."""
    message: str
    file_path: str
    canvas_display: Optional[Dict[str, Any]] = None
    can_undo: bool = False
    can_redo: bool = False

class VersionHistoryRequest(BaseModel):
    """Request to get version history."""
    file_path: str = Field(..., description="Path to the document")

class VersionHistoryResponse(BaseModel):
    """Response with version history."""
    message: str
    versions: List[Dict[str, Any]]
    current_version: int

# --- Helper Functions ---
def extract_document_content(file_path: str) -> tuple[str, str]:
    """Extract text content from various document formats."""
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(), 'text'
    
    elif file_ext == '.docx':
        from docx import Document
        doc = Document(file_path)
        content = '\n'.join([para.text for para in doc.paragraphs])
        return content, 'docx'
    
    elif file_ext == '.pdf':
        from pypdf import PdfReader
        content = []
        with open(file_path, 'rb') as f:
            pdf_reader = PdfReader(f)
            for page in pdf_reader.pages:
                content.append(page.extract_text())
        return '\n'.join(content), 'pdf'
    
    elif file_ext == '.csv':
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_markdown(index=False), 'csv'
        except ImportError:
            # Fallback to standard csv
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                # Convert to simple string or markdown-like
                header = ' | '.join(rows[0])
                separator = '|'.join(['---'] * len(rows[0]))
                body = '\n'.join([' | '.join(row) for row in rows[1:]])
                return f"{header}\n{separator}\n{body}", 'csv'
            
    elif file_ext in ['.xlsx', '.xls']:
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            return df.to_markdown(index=False), 'xlsx'
        except ImportError:
            raise ValueError(f"Pandas/Openpyxl required to read Excel files: {file_ext}")

    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def create_docx(content: str, file_path: str):
    """Create a Word document."""
    from docx import Document
    doc = Document()
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(file_path)

def create_pdf(content: str, file_path: str):
    """Create a PDF document."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for paragraph in content.split('\n'):
        if paragraph.strip():
            story.append(Paragraph(paragraph, styles['Normal']))
            story.append(Spacer(1, 12))
    
    doc.build(story)

def get_file_base64(file_path: str) -> str:
    """Convert file to base64 for display."""
    with open(file_path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8')

def create_pdf_canvas_display(file_path: str, title: str, original_type: str = 'pdf') -> Dict[str, Any]:
    """
    Create a PDF canvas display with cache busting.
    Ensures browser doesn't cache old versions.
    """
    import time
    pdf_base64 = get_file_base64(file_path)
    return create_canvas_display(
        canvas_type='pdf',
        canvas_data={
            'title': title,
            'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
            'file_path': file_path,
            'original_type': original_type,
            'timestamp': int(time.time() * 1000),  # Millisecond timestamp for cache busting
            'version': os.path.getmtime(file_path),  # File modification time
            'no_cache': True  # Hint to frontend to disable caching
        }
    )

def convert_docx_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    """
    Convert a DOCX file to PDF for display purposes.
    Returns the path to the generated PDF.
    """
    try:
        import platform
        import time
        
        if pdf_path is None:
            # Add timestamp to prevent browser caching of old PDF
            timestamp = int(time.time() * 1000)  # milliseconds
            base_name = Path(docx_path).stem
            pdf_path = str(Path(docx_path).parent / f"{base_name}_display_{timestamp}.pdf")
        
        # Use docx2pdf for Windows, LibreOffice for Linux/Mac
        if platform.system() == 'Windows':
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        else:
            # For Linux/Mac, use LibreOffice command line
            import subprocess
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', str(Path(pdf_path).parent), docx_path
            ], check=True)
            # Rename if needed
            generated_pdf = str(Path(docx_path).with_suffix('.pdf'))
            if generated_pdf != pdf_path and os.path.exists(generated_pdf):
                os.rename(generated_pdf, pdf_path)
        
        logger.info(f"Converted DOCX to PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.error(f"Failed to convert DOCX to PDF: {e}")
        raise

def analyze_document_structure(file_path: str) -> Dict[str, Any]:
    """
    Analyze the structure of a Word document to provide context for editing.
    Returns information about styles, headings, tables, images, etc.
    """
    from docx import Document
    
    doc = Document(file_path)
    
    # Collect style information
    styles_used = {}
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name not in styles_used:
            styles_used[style_name] = 0
        styles_used[style_name] += 1
    
    # Analyze headings
    headings = []
    for para in doc.paragraphs:
        if 'Heading' in para.style.name:
            headings.append({
                'level': para.style.name,
                'text': para.text[:50] + ('...' if len(para.text) > 50 else '')
            })
    
    # Count tables and images
    table_count = len(doc.tables)
    
    # Get document statistics
    total_paragraphs = len(doc.paragraphs)
    total_sections = len(doc.sections)
    
    return {
        'styles_used': styles_used,
        'headings': headings[:10],  # First 10 headings
        'table_count': table_count,
        'total_paragraphs': total_paragraphs,
        'total_sections': total_sections,
        'file_name': Path(file_path).name
    }

def interpret_editing_instruction(instruction: str, file_path: str, thread_id: str = None) -> list:
    """
    Use LLM to interpret natural language editing instruction and generate
    structured editing actions with full document context and session memory.
    """
    # Get session context for this document
    session_context = session_manager.get_session_context(file_path, thread_id)
    
    # Get all files in this conversation for cross-file awareness
    all_files_context = ""
    if thread_id:
        all_files_context = session_manager.get_all_files_summary(thread_id)
    
    # Analyze document structure
    doc_structure = analyze_document_structure(file_path)
    
    # Get current document state using intelligent editor
    try:
        intelligent_editor = IntelligentDocumentEditor(file_path)
        current_state = intelligent_editor.analyze_current_state()
    except Exception as e:
        logger.warning(f"Could not analyze document state: {e}")
        current_state = {}
    
    # Get LLM instance
    llm_instance = get_llm()
    
    # Create prompt for instruction interpretation with session context
    prompt = f"""You are a Word document editing expert with memory of previous edits. Your task is to interpret a natural language editing instruction and generate the appropriate structured editing actions.

=== ALL DOCUMENTS IN THIS CONVERSATION ===
{all_files_context if all_files_context else "Currently working with one document."}

{session_context}

CURRENT DOCUMENT STRUCTURE:
{json.dumps(doc_structure, indent=2)}

DETAILED DOCUMENT STATE:
{json.dumps(current_state, indent=2) if current_state else "Not available"}

NEW USER INSTRUCTION: "{instruction}"

AVAILABLE ACTIONS AND THEIR PARAMETERS:

**MODIFICATION ACTIONS (Use these to change existing content):**
1. convert_list_style: {{from_type: "bulleted"|"numbered", to_type: "bulleted"|"numbered", list_indices: list[int]|null}}
   - ONLY for REAL Word lists (detected in document state)
2. convert_text_to_list: {{para_indices: list[int]|null, list_type: "numbered"|"bulleted"}}
   - For plain text with list markers (-, •, 1., etc.) that need to become REAL lists
   - Auto-detects if para_indices is null
3. modify_text_formatting: {{target_style: str, target_text: str, para_indices: list[int], bold: bool, italic: bool, underline: bool, color: str, font_size: int}}
4. remove_paragraphs_by_pattern: {{pattern: str, indices: list[int]}}
5. find_and_delete: {{search_text: str, case_sensitive: bool}}
6. replace_all_text: {{old_text: str, new_text: str, case_sensitive: bool}}
7. remove_all_formatting: {{para_indices: list[int]}}
8. change_style: {{para_indices: list[int], new_style: str}}
9. delete_tables: {{table_indices: list[int]|null}}
10. modify_table: {{table_index: int, modifications: dict}}
11. remove_all_lists: {{}}
12. remove_empty_paragraphs: {{}}
13. move_paragraph: {{from_index: int, to_index: int}}
   - Move a paragraph from one position to another
14. insert_paragraph_at: {{index: int, text: str, style: str, bold: bool, italic: bool, underline: bool, color: str, font_size: int}}
   - Insert a paragraph at a specific position

**ADDITION ACTIONS (Use these to add new content):**
12. format_text: {{style: str, text: str, bold: bool, italic: bool, underline: bool, font_size: int, color: str, highlight: str}}
13. add_paragraph: {{text: str, style: str, alignment: str, bold: bool, italic: bool}}
14. add_bullet_list: {{items: list[str]}}
15. add_numbered_list: {{items: list[str]}}
16. add_heading: {{text: str, level: int}}
17. add_table: {{rows: int, cols: int, data: list[list[str]], header_row: bool}}
18. add_image: {{image_path: str, width: float, caption: str}}
19. add_page_break: {{}}
20. add_header: {{text: str, alignment: str}}
21. add_footer: {{text: str, alignment: str, page_numbers: bool}}
22. add_hyperlink: {{text: str, url: str}}

CRITICAL INTERPRETATION GUIDELINES:

**1. MODIFICATION FIRST - ADDITION LAST:**
- If user says "convert bullets to numbers" → Check document state:
  - If REAL lists exist → Use `convert_list_style`
  - If only text with markers (-, •, 1.) → Use `convert_text_to_list`
- If user says "change color" → Use `modify_text_formatting`, NOT `format_text` + `add_paragraph`
- If user says "remove extra content" → Use `find_and_delete` or `remove_paragraphs_by_pattern`
- If user says "fix numbering" or "make proper lists" → Use `convert_text_to_list`
- ONLY use addition actions when user explicitly wants NEW content added

**CRITICAL: Distinguish between REAL lists and PSEUDO-lists:**
- REAL lists: Have numbering XML, shown in document state "lists" array
- PSEUDO-lists: Plain text starting with -, •, 1., etc. (NOT in "lists" array)
- For PSEUDO-lists, ALWAYS use `convert_text_to_list`, NEVER `convert_list_style`

**2. USE SESSION CONTEXT:**
- Review the edit history to understand what was done before
- If user says "change the others too", look at previous edits to understand what "others" means
- If user says "undo that" or "remove what you added", use deletion actions
- Maintain consistency with previous edits

**3. DOCUMENT STATE AWARENESS:**
- Check CURRENT DOCUMENT STATE to see what actually exists
- "subheadings" = look at current headings in state, usually Heading 2
- "all lists" = check lists array in state
- "extra content" = compare current state with what should be there

**4. SMART INTERPRETATION:**
- Colors: Use hex codes for ANY color (e.g., "#A52A2A" for brown, "#FF0000" for red)
  - Common colors: red=#FF0000, blue=#0000FF, green=#00FF00, yellow=#FFFF00, orange=#FFA500, purple=#800080, brown=#A52A2A, pink=#FFC0CB, violet=#EE82EE
  - For any other color, look up the hex code and use it
- Ambiguous terms: use document structure to make intelligent guesses
- Multiple actions: generate as many as needed to fulfill the instruction completely

**5. DELETION AND CLEANUP:**
- If content shouldn't be there, DELETE it (don't just hide it)
- Remove empty paragraphs after deletions
- Clean up formatting when requested

**6. MOVING AND INSERTING:**
- To move content: Use `move_paragraph` with from_index and to_index
- To insert at specific location: Use `insert_paragraph_at` with index
- DON'T use find_and_delete + add_paragraph for moving (loses position control)
- Calculate indices from document state

Return a JSON array of actions. Each action should have:
{{
  "action": "action_name",
  "params": {{...parameters...}}
}}

IMPORTANT: Return ONLY the JSON array, no explanations or markdown.
"""
    
    try:
        # Get LLM response with fallback
        try:
            response = llm_instance.invoke(prompt)
        except Exception as llm_error:
            # If Cerebras fails (rate limit), try Groq as fallback
            if "429" in str(llm_error) or "rate" in str(llm_error).lower():
                logger.warning(f"Cerebras LLM failed (rate limit), trying Groq fallback")
                from langchain_groq import ChatGroq
                groq_llm = ChatGroq(model="openai/gpt-oss-120b", temperature=0)
                response = groq_llm.invoke(prompt)
            else:
                raise
        
        response_text = response.content if hasattr(response, 'content') else str(response)
        
        # Extract JSON from response
        import re
        json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
        if json_match:
            actions = json.loads(json_match.group(0))
            logger.info(f"Interpreted instruction into {len(actions)} action(s): {actions}")
            return actions
        else:
            logger.error(f"Could not extract JSON from LLM response: {response_text}")
            # Fallback: try to parse the whole response
            actions = json.loads(response_text)
            return actions
    except Exception as e:
        logger.error(f"Error interpreting instruction: {e}")
        # DO NOT add error text to document! Raise the error instead
        raise ValueError(f"Failed to interpret editing instruction: {str(e)}")

def edit_docx_advanced(file_path: str, instructions: dict, save_version: bool = True) -> str:
    """
    Advanced Word document editing using the AdvancedDocumentEditor.
    Supports comprehensive editing capabilities with version control.
    """
    # Save version before editing (for undo/redo)
    if save_version:
        try:
            action = instructions.get('action', 'edit')
            version_manager.save_version(file_path, f"Before {action}")
            logger.info(f"Saved version before editing: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to save version: {e}")
    
    # Use the advanced editor
    editor = AdvancedDocumentEditor(file_path)
    result = editor.edit(instructions)
    
    logger.info(f"Advanced edit completed: {result}")
    return file_path

# --- API Endpoints ---
@app.get("/", summary="Health Check")
def read_root():
    """Provides a simple health check endpoint."""
    return {"message": "Document Analysis Agent is running and ready to analyze, display, and edit documents."}

@app.post("/analyze",
          response_model=AnalyzeDocumentResponse,
          summary="Analyze Document(s) via RAG")
def analyze_document(request: AnalyzeDocumentRequest, thread_id: Optional[str] = None):
    """
    Analyzes one or multiple documents using a RAG pipeline. It loads pre-computed FAISS vector stores,
    retrieves relevant text chunks based on the user's query, and generates a final answer.
    
    Supports both single document (vector_store_path) and multiple documents (vector_store_paths).
    """
    # Determine which paths to use (support both single and multi-document)
    vector_store_paths_to_load = []
    if request.vector_store_paths:
        # Multiple documents provided
        # ✅ CRITICAL FIX: Filter out None values
        vector_store_paths_to_load = [p for p in request.vector_store_paths if p and p.strip()]
        if not vector_store_paths_to_load:
            raise HTTPException(
                status_code=400,
                detail="vector_store_paths contains only None or empty values. Please provide valid paths."
            )
        logger.info(f"Processing {len(vector_store_paths_to_load)} documents")
    elif request.vector_store_path:
        # Single document provided (backward compatibility)
        # ✅ CRITICAL FIX: Validate not None
        if not request.vector_store_path or request.vector_store_path.strip() == '':
            raise HTTPException(
                status_code=400,
                detail="vector_store_path is None or empty. Please provide a valid path."
            )
        vector_store_paths_to_load = [request.vector_store_path]
        logger.info(f"Processing single document")
    else:
        raise HTTPException(
            status_code=400,
            detail="Either vector_store_path or vector_store_paths must be provided"
        )
    
    # Validate that all vector stores exist
    for vsp in vector_store_paths_to_load:
        if not os.path.exists(vsp):
            logger.error(f"Document vector store not found at path: {vsp}")
            raise HTTPException(
                status_code=404,
                detail=f"Document vector store not found at path: {vsp}"
            )

    try:
        # Lazy import heavy dependencies
        from langchain_community.vectorstores import FAISS
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import StrOutputParser
        from langchain_core.runnables import RunnablePassthrough
        
        # Get the lazy-loaded embeddings and LLM
        embeddings = get_embeddings()
        llm_instance = get_llm()
        
        # Load and merge all vector stores
        logger.info(f"Loading {len(vector_store_paths_to_load)} vector store(s)")
        combined_vector_store = None
        
        for idx, vsp in enumerate(vector_store_paths_to_load):
            logger.info(f"Loading vector store {idx+1}/{len(vector_store_paths_to_load)}: {os.path.basename(vsp)}")
            vs = FAISS.load_local(
                vsp,
                embeddings,
                allow_dangerous_deserialization=True
            )
            
            if combined_vector_store is None:
                combined_vector_store = vs
            else:
                # Merge this vector store into the combined one
                combined_vector_store.merge_from(vs)
                logger.info(f"Merged vector store {idx+1} into combined store")
        
        logger.info(f"Combined vector store ready with documents from {len(vector_store_paths_to_load)} files")

        logger.info(f"Combined vector store ready with documents from {len(vector_store_paths_to_load)} files")

        # 3. Create a modern RAG chain using LCEL (LangChain Expression Language)
        retriever = combined_vector_store.as_retriever(search_kwargs={"k": 5})
        
        # Define the prompt template for RAG
        template = """Answer the question based only on the following context:

{context}

Question: {question}

Answer:"""
        
        prompt = ChatPromptTemplate.from_template(template)
        
        # Helper function to format retrieved documents
        def format_docs(docs):
            return "\n\n".join(doc.page_content for doc in docs)
        
        # Build the RAG chain using LCEL
        rag_chain = (
            {"context": retriever | format_docs, "question": RunnablePassthrough()}
            | prompt
            | llm_instance
            | StrOutputParser()
        )

        # 4. Invoke the chain to get the answer. This performs the retrieval and generation steps.
        logger.info(f"Processing query: {request.query}")
        result = rag_chain.invoke(request.query)

        # 5. Return the generated answer in the specified response format.
        logger.info("Document analysis completed successfully")
        
        # 6. Create canvas display to show the original document
        # Only attempt for single-document mode to avoid NoneType errors
        canvas_display = None
        
        # CRITICAL FIX: Only create canvas if vector_store_path exists (single-document mode)
        # In multi-document mode, vector_store_path is None and vector_store_paths is used
        if request.vector_store_path:
            try:
                # Vector store path is typically: storage/vector_store/{doc_name}.pdf.faiss
                vector_store_base = Path(request.vector_store_path).stem  # e.g., "TRM.pdf"
                
                # Try to find the document in storage/documents
                possible_doc_paths = [
                    f"storage/documents/{vector_store_base}",
                    f"storage/documents\\{vector_store_base}",  # Windows path
                ]
                
                doc_path = None
                for path in possible_doc_paths:
                    if os.path.exists(path):
                        doc_path = path
                        break
                
                if doc_path and doc_path.lower().endswith('.pdf'):
                    # For PDFs, show the actual PDF in canvas
                    logger.info(f"Creating PDF canvas display for: {doc_path}")
                    try:
                        pdf_base64 = get_file_base64(doc_path)
                        canvas_display = create_canvas_display(
                            canvas_type='pdf',
                            canvas_data={
                                'title': Path(doc_path).name,
                                'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                                'file_path': doc_path
                            }
                        )
                    except Exception as pdf_err:
                        logger.error(f"Failed to create PDF canvas: {pdf_err}")
                elif doc_path and doc_path.lower().endswith('.docx'):
                    # For Word documents, convert to PDF for display
                    logger.info(f"Creating PDF canvas display for Word doc: {doc_path}")
                    try:
                        pdf_path = convert_docx_to_pdf(doc_path)
                        pdf_base64 = get_file_base64(pdf_path)
                        canvas_display = create_canvas_display(
                            canvas_type='pdf',
                            canvas_data={
                                'title': Path(doc_path).name,
                                'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                                'file_path': doc_path,
                                'original_type': 'docx'
                            }
                        )
                    except Exception as docx_err:
                        logger.error(f"Failed to convert DOCX to PDF: {docx_err}")
            except Exception as canvas_err:
                # Non-fatal: canvas creation failed but analysis succeeded
                logger.warning(f"Canvas display creation failed (non-fatal): {canvas_err}")
                canvas_display = None
        else:
            # Multi-document mode - no single canvas display
            # The analysis answer contains consolidated information from all documents
            logger.info("Multi-document mode: skipping canvas display, answer contains consolidated results")
        
        return AnalyzeDocumentResponse(answer=result, canvas_display=canvas_display)

    except Exception as e:
        # Catch-all for any other errors during the process
        logger.error(f"Error occurred while analyzing the document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred while analyzing the document: {str(e)}"
        )

@app.post("/display",
          response_model=DisplayDocumentResponse,
          summary="Display a Document in Canvas")
def display_document(request: DisplayDocumentRequest, thread_id: Optional[str] = None):
    """
    Displays a document (PDF, DOCX, TXT) in the canvas by extracting its content
    and formatting it for display.
    
    Parameters:
        request: The display request with file path
        thread_id: Optional conversation thread ID for context
    """
    # Normalize path to OS-specific format (handles both forward and backward slashes)
    normalized_path = os.path.normpath(request.file_path)
    
    # If path doesn't exist, try prepending parent directory (for paths relative to backend/)
    if not os.path.exists(normalized_path):
        # Try path relative to parent directory (backend/)
        parent_path = os.path.join('..', normalized_path)
        if os.path.exists(parent_path):
            normalized_path = parent_path
            logger.info(f"📄 Display request: Found file in parent directory: {normalized_path}")
    
    logger.info(f"📄 Display request received: original_path={request.file_path}, normalized_path={normalized_path}")
    
    if not os.path.exists(normalized_path):
        logger.error(f"Document not found at path: {normalized_path} (original: {request.file_path})")
        raise HTTPException(
            status_code=404,
            detail=f"Document not found at path: {request.file_path}"
        )
    
    # Use normalized path for all operations
    request.file_path = normalized_path
    
    try:
        # Extract content from document
        content, doc_type = extract_document_content(request.file_path)
        file_name = Path(request.file_path).name
        
        # Create canvas display based on document type
        if doc_type == 'pdf':
            # For PDFs, provide base64 encoded file for PDF viewer
            pdf_base64 = get_file_base64(request.file_path)
            canvas_display = create_canvas_display(
                canvas_type='pdf',
                canvas_data={
                    'title': file_name,
                    'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                    'file_path': request.file_path
                }
            )
        elif doc_type == 'docx':
            # For Word documents, convert to PDF for display (preserves formatting)
            logger.info(f"🔄 Starting DOCX to PDF conversion for: {file_name}")
            try:
                logger.info(f"🔄 Calling convert_docx_to_pdf({request.file_path})")
                pdf_path = convert_docx_to_pdf(request.file_path)
                logger.info(f"🔄 PDF created at: {pdf_path}")
                pdf_base64 = get_file_base64(pdf_path)
                logger.info(f"🔄 PDF encoded to base64, length: {len(pdf_base64)}")
                canvas_display = create_canvas_display(
                    canvas_type='pdf',
                    canvas_data={
                        'title': file_name,
                        'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                        'file_path': request.file_path,
                        'original_type': 'docx'
                    }
                )
                logger.info(f"✅ Converted DOCX to PDF for display: {file_name}")
            except Exception as conv_err:
                logger.error(f"❌ Failed to convert DOCX to PDF, showing text: {conv_err}", exc_info=True)
                # Fallback to text display
                canvas_display = create_canvas_display(
                    canvas_type='document',
                    canvas_data={
                        'title': file_name,
                        'content': content,
                        'file_path': request.file_path,
                        'file_type': doc_type
                    }
                )
        else:
            # For text-based documents, show formatted content
            canvas_display = create_canvas_display(
                canvas_type='document',
                canvas_data={
                    'title': file_name,
                    'content': content,
                    'file_path': request.file_path,
                    'file_type': doc_type
                }
            )
        
        logger.info(f"✅ Document displayed successfully: {file_name}")
        logger.info(f"📄 Returning canvas display: type={canvas_display.get('canvas_type')}, title={canvas_display.get('canvas_title')}")
        return DisplayDocumentResponse(
            message=f"Document '{file_name}' displayed successfully",
            canvas_display=canvas_display,
            file_type=doc_type
        )
        
    except Exception as e:
        logger.error(f"Error displaying document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to display document: {str(e)}"
        )

@app.post("/create",
          response_model=CreateDocumentResponse,
          summary="Create a New Document")
def create_document(request: CreateDocumentRequest):
    """
    Creates a new document (DOCX, TXT, or PDF) with optional LLM-powered content generation.
    
    Supports:
    1. Direct content: Just provide content and create the document
    2. Reference file: Provide reference_file_path to transform an existing file
    3. LLM generation: Provide instruction to have LLM generate content
    
    Examples:
    - {"content": "My report content...", "file_name": "report.docx"}
    - {"reference_file_path": "storage/documents/data.pdf", "instruction": "Create a summary report", "file_name": "summary.docx"}
    - {"instruction": "Create a professional cover letter for a software engineer", "file_name": "cover_letter.docx"}
    """
    try:
        # Ensure output directory exists
        os.makedirs(request.output_dir, exist_ok=True)
        
        # Build full file path
        file_path = os.path.join(request.output_dir, request.file_name)
        
        # Step 1: Determine content source
        final_content = None
        
        if request.content:
            # Direct content provided
            final_content = request.content
            
            # If instruction is also provided, use LLM to transform content
            if request.instruction:
                logger.info(f"Using LLM to transform provided content with instruction: {request.instruction}")
                final_content = _generate_document_content_with_llm(request.content, request.instruction)
                if not final_content:
                    final_content = request.content  # Fallback to original
                    
        elif request.reference_file_path:
            # Load content from reference file
            if os.path.exists(request.reference_file_path):
                source_content, source_type = extract_document_content(request.reference_file_path)
                logger.info(f"Loaded source content from: {request.reference_file_path} ({len(source_content)} chars)")
                
                if request.instruction:
                    logger.info(f"Using LLM to transform source content with instruction: {request.instruction}")
                    final_content = _generate_document_content_with_llm(source_content, request.instruction)
                else:
                    final_content = source_content
            else:
                raise ValueError(f"Reference file not found: {request.reference_file_path}")
                
        elif request.instruction:
            # Generate from scratch with LLM
            logger.info(f"Using LLM to generate document from instruction: {request.instruction}")
            final_content = _generate_document_content_with_llm(None, request.instruction)
            if not final_content:
                raise ValueError("Failed to generate content with LLM")
        else:
            raise ValueError("Either 'content', 'reference_file_path', or 'instruction' must be provided")
        
        # Step 2: Create the document
        if request.file_type == 'txt':
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(final_content)
        elif request.file_type == 'docx':
            create_docx(final_content, file_path)
        elif request.file_type == 'pdf':
            create_pdf(final_content, file_path)
        else:
            raise ValueError(f"Unsupported file type: {request.file_type}")
        
        logger.info(f"Document created successfully: {file_path}")
        
        # Step 3: Create canvas display
        if request.file_type == 'docx':
            try:
                pdf_path = convert_docx_to_pdf(file_path)
                pdf_base64 = get_file_base64(pdf_path)
                canvas_display = create_canvas_display(
                    canvas_type='pdf',
                    canvas_data={
                        'title': request.file_name,
                        'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                        'file_path': file_path,
                        'original_type': 'docx'
                    }
                )
            except Exception as conv_err:
                logger.warning(f"Failed to convert created DOCX to PDF: {conv_err}")
                canvas_display = create_canvas_display(
                    canvas_type='document',
                    canvas_data={
                        'title': request.file_name,
                        'content': final_content,
                        'file_path': file_path,
                        'file_type': request.file_type
                    }
                )
        else:
            canvas_display = create_canvas_display(
                canvas_type='document',
                canvas_data={
                    'title': request.file_name,
                    'content': final_content,
                    'file_path': file_path,
                    'file_type': request.file_type
                }
            )
        
        return CreateDocumentResponse(
            message=f"✅ Document '{request.file_name}' created successfully",
            file_path=file_path,
            canvas_display=canvas_display
        )
        
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create document: {str(e)}"
        )


def _generate_document_content_with_llm(source_content: Optional[str], instruction: str) -> Optional[str]:
    """Use LLM to generate or transform document content."""
    llm_instance = get_llm()
    
    if source_content:
        prompt = f"""You are a professional document writer. Transform the following content based on the instruction.

=== SOURCE CONTENT ===
{source_content[:8000]}

=== INSTRUCTION ===
{instruction}

=== RULES ===
1. Output ONLY the final document content - no explanations
2. Maintain professional formatting
3. Use appropriate paragraphs and structure
4. If creating a structured document, use clear headings

=== OUTPUT DOCUMENT CONTENT NOW ==="""
    else:
        prompt = f"""You are a professional document writer. Create a document based on the following instruction.

=== INSTRUCTION ===
{instruction}

=== RULES ===
1. Output ONLY the document content - no explanations or meta-commentary
2. Use professional language and formatting
3. Structure with clear paragraphs
4. Include appropriate headings if needed

=== OUTPUT DOCUMENT CONTENT NOW ==="""
    
    try:
        response = llm_instance.invoke(prompt)
        content = response.content if hasattr(response, 'content') else str(response)
        logger.info(f"LLM generated document content: {len(content)} chars")
        return content.strip()
    except Exception as e:
        logger.error(f"LLM document generation failed: {e}")
        return None


@app.post("/edit",
          response_model=EditDocumentResponse,
          summary="Edit an Existing Document with Natural Language")
def edit_document(request: EditDocumentRequest, thread_id: Optional[str] = None):
    """
    Edits an existing document using natural language instruction.
    The agent interprets the instruction with full document context and executes appropriate edits.
    Always executes the edit immediately.
    
    Parameters:
        request: The edit request with file path and instruction
        thread_id: Optional conversation thread ID for isolation across conversations
    """
    # ✅ CRITICAL FIX: Validate file_path is not None
    if not request.file_path or request.file_path.strip() == '':
        logger.error("❌ Edit request failed: file_path is None or empty")
        raise HTTPException(
            status_code=400,
            detail="file_path is required and cannot be None or empty. Please provide a valid document path."
        )
    
    # Normalize path to OS-specific format (handles both forward and backward slashes)
    normalized_path = os.path.normpath(request.file_path)
    
    # If path doesn't exist, try prepending parent directory (for paths relative to backend/)
    if not os.path.exists(normalized_path):
        # Try path relative to parent directory (backend/)
        parent_path = os.path.join('..', normalized_path)
        if os.path.exists(parent_path):
            normalized_path = parent_path
            logger.info(f"📄 Edit request: Found file in parent directory: {normalized_path}")
    
    logger.info(f"📄 Edit request received: original_path={request.file_path}, normalized_path={normalized_path}")
    logger.info(f"📝 Natural language instruction: {request.instruction}")
    
    if not os.path.exists(normalized_path):
        logger.error(f"Document not found at path: {normalized_path} (original: {request.file_path})")
        raise HTTPException(
            status_code=404,
            detail=f"Document not found at path: {request.file_path}"
        )
    
    # Use normalized path for all operations
    request.file_path = normalized_path
    
    try:
        file_name = Path(request.file_path).name
        file_ext = Path(request.file_path).suffix.lower()
        
        # Execute the edit based on file type
        if file_ext == '.docx':
            # Get document state BEFORE editing
            intelligent_editor = IntelligentDocumentEditor(request.file_path)
            state_before = intelligent_editor.analyze_current_state()
            
            # Interpret natural language instruction into structured actions (with session context)
            logger.info(f"🧠 Interpreting instruction with document context and session memory...")
            actions = interpret_editing_instruction(request.instruction, request.file_path)
            logger.info(f"✅ Generated {len(actions)} editing action(s): {[a.get('action') for a in actions]}")
            
            # Save version before editing
            try:
                version_manager.save_version(request.file_path, f"Before: {request.instruction[:50]}")
                logger.info(f"💾 Saved version before editing")
            except Exception as e:
                logger.warning(f"Failed to save version: {e}")
            
            # Execute actions using BOTH editors (intelligent for modifications, advanced for additions)
            results = []
            for action in actions:
                action_name = action.get('action')
                params = action.get('params', {})
                
                try:
                    # Route to appropriate editor based on action type
                    if action_name in ['convert_list_style', 'modify_text_formatting', 'remove_paragraphs_by_pattern',
                                      'find_and_delete', 'replace_all_text', 'remove_all_formatting', 'change_style',
                                      'delete_tables', 'modify_table', 'remove_all_lists', 'remove_empty_paragraphs',
                                      'convert_text_to_list', 'detect_pseudo_lists', 'move_paragraph', 'insert_paragraph_at']:
                        # Use intelligent editor for modification actions
                        if action_name == 'convert_list_style':
                            intelligent_editor.convert_list_style(
                                params.get('from_type'), 
                                params.get('to_type'),
                                params.get('list_indices')
                            )
                            result = f"Converted lists from {params.get('from_type')} to {params.get('to_type')}"
                        elif action_name == 'modify_text_formatting':
                            intelligent_editor.modify_text_formatting(**params)
                            result = f"Modified text formatting"
                        elif action_name == 'remove_paragraphs_by_pattern':
                            intelligent_editor.remove_paragraphs_by_pattern(**params)
                            result = f"Removed paragraphs by pattern"
                        elif action_name == 'find_and_delete':
                            count = intelligent_editor.find_and_delete(**params)
                            result = f"Deleted {count} paragraphs"
                        elif action_name == 'replace_all_text':
                            count = intelligent_editor.replace_all_text(**params)
                            result = f"Replaced {count} occurrences"
                        elif action_name == 'remove_all_formatting':
                            intelligent_editor.remove_all_formatting(**params)
                            result = f"Removed all formatting"
                        elif action_name == 'change_style':
                            intelligent_editor.change_style(**params)
                            result = f"Changed paragraph styles"
                        elif action_name == 'delete_tables':
                            intelligent_editor.delete_tables(**params)
                            result = f"Deleted tables"
                        elif action_name == 'modify_table':
                            intelligent_editor.modify_table(**params)
                            result = f"Modified table"
                        elif action_name == 'remove_all_lists':
                            intelligent_editor.remove_all_lists()
                            result = f"Removed all list formatting"
                        elif action_name == 'remove_empty_paragraphs':
                            count = intelligent_editor.remove_empty_paragraphs()
                            result = f"Removed {count} empty paragraphs"
                        elif action_name == 'convert_text_to_list':
                            count = intelligent_editor.convert_text_to_list(**params)
                            result = f"Converted {count} text paragraphs to Word lists"
                        elif action_name == 'detect_pseudo_lists':
                            pseudo_lists = intelligent_editor.detect_pseudo_lists()
                            result = f"Detected {len(pseudo_lists)} pseudo-list items"
                        elif action_name == 'move_paragraph':
                            count = intelligent_editor.move_paragraph(**params)
                            result = f"Moved {count} paragraph"
                        elif action_name == 'insert_paragraph_at':
                            count = intelligent_editor.insert_paragraph_at(**params)
                            result = f"Inserted {count} paragraph"
                        
                        # Save after intelligent editor actions
                        intelligent_editor.save()
                        logger.info(f"✅ Executed intelligent action '{action_name}': {result}")
                    else:
                        # Use advanced editor for addition actions
                        editor = AdvancedDocumentEditor(request.file_path)
                        result = editor.edit(action)
                        logger.info(f"✅ Executed advanced action '{action_name}': {result}")
                    
                    results.append(result)
                except Exception as e:
                    logger.error(f"❌ Failed to execute action '{action_name}': {e}")
                    results.append(f"Error in {action_name}: {str(e)}")
            
            # Get document state AFTER editing
            intelligent_editor = IntelligentDocumentEditor(request.file_path)  # Reload
            state_after = intelligent_editor.analyze_current_state()
            
            # VALIDATE STATE CHANGE: Verify that the edit actually modified the document
            # This addresses the unreliable editing issue by confirming changes were applied
            state_changed = state_before != state_after
            
            if not state_changed:
                logger.warning(f"⚠️  STATE VALIDATION WARNING: Document state did not change after edit.")
                logger.warning(f"   Actions executed: {[a.get('action') for a in actions]}")
                logger.warning(f"   Before: {state_before}")
                logger.warning(f"   After: {state_after}")
                # Add validation note to results
                results.append("⚠️  Warning: Document state did not appear to change. Edit may not have been applied correctly.")
            else:
                logger.info(f"✅ STATE VALIDATION PASSED: Document state changed as expected after edit")
            
# Record edit in session with thread isolation
            session_manager.add_edit_action(
                document_path=request.file_path,
                action_type="edit",
                instruction=request.instruction,
                parameters={'actions': actions},
                result="; ".join(results),
                state_before=state_before,
                state_after=state_after,
                thread_id=thread_id or "default"
            )
            
            # Record conversation turn with thread isolation
            session_manager.add_conversation_turn(
                document_path=request.file_path,
                user_message=request.instruction,
                agent_response="; ".join(results),
                thread_id=thread_id or "default"
            )
            
            logger.info(f"✅ All actions executed and recorded in session: {results}")
            
        elif file_ext == '.txt':
            # For text files, treat instruction as new content
            # (or we could use LLM to interpret, but keeping it simple for now)
            with open(request.file_path, 'w', encoding='utf-8') as f:
                f.write(request.instruction)
            logger.info(f"✅ Text file updated")
            
        elif file_ext == '.pdf':
            # PDFs are harder to edit, so we'll just note the limitation
            raise ValueError("PDF editing with natural language is not yet supported. Please use DOCX format.")
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        logger.info(f"✅ Document edited successfully: {request.file_path}")
        
        # Create canvas display showing the edited document
        # For DOCX files, convert to PDF for display (preserves formatting)
        if file_ext == '.docx':
            logger.info(f"🔄 Converting edited DOCX to PDF for display: {file_name}")
            try:
                pdf_path = convert_docx_to_pdf(request.file_path)
                pdf_base64 = get_file_base64(pdf_path)
                canvas_display = create_canvas_display(
                    canvas_type='pdf',
                    canvas_data={
                        'title': file_name,
                        'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                        'file_path': request.file_path,
                        'original_type': 'docx'
                    }
                )
                logger.info(f"✅ Edited DOCX converted to PDF for display: {file_name}")
            except Exception as conv_err:
                logger.warning(f"Failed to convert edited DOCX to PDF, showing text: {conv_err}")
                # Fallback to text display
                content, _ = extract_document_content(request.file_path)
                canvas_display = create_canvas_display(
                    canvas_type='document',
                    canvas_data={
                        'title': file_name,
                        'content': content,
                        'file_path': request.file_path
                    }
                )
        else:
            # For other file types, show text content
            content = request.content if request.content else extract_document_content(request.file_path)[0]
            canvas_display = create_canvas_display(
                canvas_type='document',
                canvas_data={
                    'title': file_name,
                    'content': content,
                    'file_path': request.file_path
                }
            )
        
        logger.info(f"📄 Returning canvas display for edited document: type={canvas_display.get('canvas_type')}")
        
        # Check undo/redo availability
        can_undo = version_manager.can_undo(request.file_path)
        can_redo = version_manager.can_redo(request.file_path)
        
        return EditDocumentResponse(
            message=f"✅ Document '{file_name}' edited successfully",
            file_path=request.file_path,
            canvas_display=canvas_display,
            can_undo=can_undo,
            can_redo=can_redo
        )
        
    except Exception as e:
        logger.error(f"Error editing document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to edit document: {str(e)}"
        )

@app.post("/undo",
          response_model=UndoRedoResponse,
          summary="Undo Last Edit")
def undo_edit(request: UndoRedoRequest):
    """
    Undo the last edit to a document.
    """
    # Normalize path
    normalized_path = os.path.normpath(request.file_path)
    if not os.path.exists(normalized_path):
        parent_path = os.path.join('..', normalized_path)
        if os.path.exists(parent_path):
            normalized_path = parent_path
    
    if not os.path.exists(normalized_path):
        raise HTTPException(
            status_code=404,
            detail=f"Document not found at path: {request.file_path}"
        )
    
    try:
        # Perform undo
        version_info = version_manager.undo(normalized_path)
        
        if not version_info:
            raise HTTPException(
                status_code=400,
                detail="No previous version available to undo"
            )
        
        file_name = Path(normalized_path).name
        file_ext = Path(normalized_path).suffix.lower()
        
        # Create canvas display for the restored version
        if file_ext == '.docx':
            try:
                pdf_path = convert_docx_to_pdf(normalized_path)
                pdf_base64 = get_file_base64(pdf_path)
                canvas_display = create_canvas_display(
                    canvas_type='pdf',
                    canvas_data={
                        'title': file_name,
                        'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                        'file_path': normalized_path,
                        'original_type': 'docx'
                    }
                )
            except Exception as conv_err:
                logger.warning(f"Failed to convert DOCX to PDF: {conv_err}")
                content, _ = extract_document_content(normalized_path)
                canvas_display = create_canvas_display(
                    canvas_type='document',
                    canvas_data={
                        'title': file_name,
                        'content': content,
                        'file_path': normalized_path
                    }
                )
        else:
            content, _ = extract_document_content(normalized_path)
            canvas_display = create_canvas_display(
                canvas_type='document',
                canvas_data={
                    'title': file_name,
                    'content': content,
                    'file_path': normalized_path
                }
            )
        
        can_undo = version_manager.can_undo(normalized_path)
        can_redo = version_manager.can_redo(normalized_path)
        
        return UndoRedoResponse(
            message=f"⏪ Undone to previous version: {version_info['description']}",
            file_path=normalized_path,
            canvas_display=canvas_display,
            can_undo=can_undo,
            can_redo=can_redo
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error undoing edit: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to undo edit: {str(e)}"
        )

@app.post("/redo",
          response_model=UndoRedoResponse,
          summary="Redo Last Undone Edit")
def redo_edit(request: UndoRedoRequest):
    """
    Redo the last undone edit to a document.
    """
    # Normalize path
    normalized_path = os.path.normpath(request.file_path)
    if not os.path.exists(normalized_path):
        parent_path = os.path.join('..', normalized_path)
        if os.path.exists(parent_path):
            normalized_path = parent_path
    
    if not os.path.exists(normalized_path):
        raise HTTPException(
            status_code=404,
            detail=f"Document not found at path: {request.file_path}"
        )
    
    try:
        # Perform redo
        version_info = version_manager.redo(normalized_path)
        
        if not version_info:
            raise HTTPException(
                status_code=400,
                detail="No next version available to redo"
            )
        
        file_name = Path(normalized_path).name
        file_ext = Path(normalized_path).suffix.lower()
        
        # Create canvas display for the restored version
        if file_ext == '.docx':
            try:
                pdf_path = convert_docx_to_pdf(normalized_path)
                pdf_base64 = get_file_base64(pdf_path)
                canvas_display = create_canvas_display(
                    canvas_type='pdf',
                    canvas_data={
                        'title': file_name,
                        'pdf_data': f'data:application/pdf;base64,{pdf_base64}',
                        'file_path': normalized_path,
                        'original_type': 'docx'
                    }
                )
            except Exception as conv_err:
                logger.warning(f"Failed to convert DOCX to PDF: {conv_err}")
                content, _ = extract_document_content(normalized_path)
                canvas_display = create_canvas_display(
                    canvas_type='document',
                    canvas_data={
                        'title': file_name,
                        'content': content,
                        'file_path': normalized_path
                    }
                )
        else:
            content, _ = extract_document_content(normalized_path)
            canvas_display = create_canvas_display(
                canvas_type='document',
                canvas_data={
                    'title': file_name,
                    'content': content,
                    'file_path': normalized_path
                }
            )
        
        can_undo = version_manager.can_undo(normalized_path)
        can_redo = version_manager.can_redo(normalized_path)
        
        return UndoRedoResponse(
            message=f"⏩ Redone to next version: {version_info['description']}",
            file_path=normalized_path,
            canvas_display=canvas_display,
            can_undo=can_undo,
            can_redo=can_redo
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error redoing edit: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to redo edit: {str(e)}"
        )

@app.post("/history",
          response_model=VersionHistoryResponse,
          summary="Get Document Version History")
def get_version_history(request: VersionHistoryRequest):
    """
    Get the version history for a document.
    """
    # Normalize path
    normalized_path = os.path.normpath(request.file_path)
    if not os.path.exists(normalized_path):
        parent_path = os.path.join('..', normalized_path)
        if os.path.exists(parent_path):
            normalized_path = parent_path
    
    if not os.path.exists(normalized_path):
        raise HTTPException(
            status_code=404,
            detail=f"Document not found at path: {request.file_path}"
        )
    
    try:
        versions = version_manager.get_history(normalized_path)
        doc_key = version_manager._get_document_key(normalized_path)
        current_version = version_manager.index.get(doc_key, {}).get("current_version", -1)
        
        return VersionHistoryResponse(
            message=f"Found {len(versions)} versions",
            versions=versions,
            current_version=current_version
        )
        
    except Exception as e:
        logger.error(f"Error getting version history: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get version history: {str(e)}"
        )

if __name__ == "__main__":
    # This block allows you to run the server directly for testing
    # Use: python backend/agents/document_analysis_agent.py
    port = int(os.getenv("DOCUMENT_AGENT_PORT", 8070))
    logger.info(f"Starting Document Analysis Agent on port {port}")
    # reload=False for faster startup - use reload=True only during development
    uvicorn.run("document_analysis_agent:app", host="0.0.0.0", port=port, reload=False)
